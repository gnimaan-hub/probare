"""Génère le document de documentation de Probare."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Styles de base ─────────────────────────────────────────────────────────────
styles = doc.styles

def set_font(run, size=11, bold=False, italic=False, color=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x3B, 0x30, 0xA8)  # indigo
    p.runs[0].font.size = Pt(18)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)  # violet
    p.runs[0].font.size = Pt(14)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)  # slate
    p.runs[0].font.size = Pt(12)
    return p

def body(text, bold_parts=None):
    """Ajoute un paragraphe. bold_parts = liste de sous-chaînes à mettre en gras."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_parts:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                r = p.add_run(remaining[:idx])
                r.font.size = Pt(11)
            r = p.add_run(bp)
            r.font.size = Pt(11)
            r.bold = True
            remaining = remaining[idx+len(bp):]
        if remaining:
            r = p.add_run(remaining)
            r.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p

def fichier_ref(text):
    """Ajoute un bloc de référence de fichier stylisé."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("📁 Fichier : ")
    r.font.size = Pt(10)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    r2.font.name = "Courier New"
    return p

def code_inline(p, text):
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
    return r

def note_box(text, color=(0xEF, 0xF6, 0xFF)):
    """Paragraphe encadré visuellement."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    r = p.add_run(text)
    r.font.size = Pt(11)
    return p

def separator():
    doc.add_paragraph("─" * 80)

def spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)


# ══════════════════════════════════════════════════════════════════════════════
# PAGE DE TITRE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\n\n\n")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PROBARE")
r.font.size = Pt(40)
r.font.bold = True
r.font.color.rgb = RGBColor(0x3B, 0x30, 0xA8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Logiciel d'Audit Comptable Assisté")
r.font.size = Pt(20)
r.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Documentation technique et fonctionnelle — MVP v0.1")
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("\nCabinet d'audit · Juin 2026")
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SOMMAIRE MANUEL
# ══════════════════════════════════════════════════════════════════════════════
heading1("Sommaire")
toc_items = [
    ("1.", "Qu'est-ce que Probare ?"),
    ("2.", "Architecture globale"),
    ("3.", "Les 7 étapes du pipeline d'audit"),
    ("4.", "Le moteur Python — apps/engine"),
    ("  4.1", "La provenance des données (DonneeSourcee)"),
    ("  4.2", "L'ingestion des fichiers comptables"),
    ("  4.3", "Le registre des contrôles NEP"),
    ("  4.4", "Les contrôles déterministes"),
    ("  4.5", "La machine à états"),
    ("  4.6", "La base de données SQLite"),
    ("  4.7", "L'anonymisation pour les appels IA"),
    ("  4.8", "L'interface LLM"),
    ("  4.9", "Les exports de livrables"),
    ("  4.10", "L'API REST (FastAPI)"),
    ("5.", "L'interface desktop — apps/desktop"),
    ("  5.1", "Le processus principal Electron"),
    ("  5.2", "Le sidecar Python"),
    ("  5.3", "L'interface React"),
    ("  5.4", "Les pages une par une"),
    ("6.", "Le flux complet d'une mission d'audit"),
    ("7.", "Les règles non négociables"),
    ("8.", "Glossaire"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(num + "  ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    r2 = p.add_run(title)
    r2.font.size = Pt(11)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 1. QU'EST-CE QUE PROBARE ?
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Qu'est-ce que Probare ?")

body(
    "Probare est un logiciel de bureau conçu pour assister les auditeurs légaux dans la réalisation "
    "de leurs missions d'audit comptable. Son nom vient du latin « prouver » — ce qui résume "
    "parfaitement sa philosophie : chaque chiffre produit par le logiciel doit être prouvable, "
    "traçable, et vérifiable jusqu'à sa source originale.",
    bold_parts=["Probare", "chaque chiffre produit par le logiciel doit être prouvable, traçable, et vérifiable"]
)

body(
    "L'audit légal est un processus réglementé. En France, les auditeurs (commissaires aux comptes) "
    "doivent suivre des Normes d'Exercice Professionnel (NEP) qui définissent précisément comment "
    "examiner les comptes d'une entreprise. Ces normes imposent notamment de :"
)
bullet("Justifier chaque montant contrôlé par une référence à un document source (NEP 500 — Éléments probants)")
bullet("Documenter toutes les anomalies détectées et les décisions prises (NEP 230 — Documentation)")
bullet("Identifier les zones de risque et adapter les contrôles en conséquence (NEP 330 — Procédures d'audit)")
bullet("Analyser les variations inhabituelles entre deux exercices (NEP 520 — Procédures analytiques)")

spacer()
body(
    "Probare automatise la partie mécanique de ces contrôles (les calculs, les vérifications d'équilibre, "
    "la détection d'anomalies) tout en laissant à l'auditeur humain la responsabilité des jugements "
    "professionnels. Le logiciel ne signe pas, ne décide pas — il assiste.",
    bold_parts=["laissant à l'auditeur humain la responsabilité des jugements professionnels", "ne signe pas, ne décide pas"]
)

note_box(
    "💡 Analogie : Probare est à l'auditeur ce qu'un tableur expert automatisé serait à un comptable — "
    "il fait les calculs fastidieux, conserve les preuves, et présente les résultats. "
    "L'humain interprète, décide, et signe."
)

heading2("Ce que Probare fait")
bullet("Importe et analyse des fichiers comptables (balance, grand livre) en Excel ou CSV")
bullet("Exécute 10 contrôles déterministes conformes aux NEP (calculs purement en code Python)")
bullet("Identifie et classe les anomalies en exceptions (mineures, significatives, critiques)")
bullet("Permet à l'auditeur de documenter sa décision sur chaque exception (tranchement)")
bullet("Génère un dossier de travail (.docx) et un tableau des exceptions (.xlsx)")
bullet("Conserve une piste d'audit complète de toutes les actions effectuées")
bullet("Propose (optionnellement) une assistance IA (Claude) pour l'interprétation des anomalies")

heading2("Ce que Probare ne fait jamais")
bullet("Probare ne calcule jamais via l'IA — tout calcul est du code Python déterministe")
bullet("Probare ne stocke jamais la clé API dans la base de données ou dans le code")
bullet("Probare ne transmet jamais les vraies identités clients à un service externe (anonymisation)")
bullet("Probare ne permet pas à l'IA de prendre des décisions d'audit — c'est toujours l'auditeur")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 2. ARCHITECTURE GLOBALE
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Architecture globale")

body(
    "Probare est structuré comme un monorepo (un seul dépôt contenant plusieurs applications). "
    "Il se décompose en deux grandes parties qui communiquent entre elles :"
)

heading2("2.1 L'application de bureau (apps/desktop)")
body(
    "C'est ce que l'utilisateur voit et utilise. Elle est construite avec Electron, une technologie "
    "qui permet de créer des applications de bureau en utilisant les technologies web (HTML, CSS, "
    "JavaScript). Concrètement, c'est une fenêtre de navigateur déguisée en application native."
)
body(
    "L'interface est développée avec React (bibliothèque de composants d'interface), TypeScript "
    "(JavaScript avec typage fort pour éviter les bugs), et Tailwind CSS (système de mise en page). "
    "Les animations fluides utilisent Framer Motion."
)
fichier_ref("apps/desktop/src/renderer/src/  — tout le code de l'interface")
fichier_ref("apps/desktop/src/main/  — le processus principal Electron")

heading2("2.2 Le moteur de calcul (apps/engine)")
body(
    "C'est le cerveau caché du logiciel. C'est une application Python qui tourne en arrière-plan "
    "sur l'ordinateur de l'auditeur. Elle expose une API REST (un ensemble d'URLs qu'on peut "
    "appeler pour demander des calculs, enregistrer des données, etc.) via FastAPI."
)
body(
    "Cette architecture est appelée « sidecar » — le moteur Python est lancé automatiquement "
    "par l'application de bureau au démarrage, et arrêté à la fermeture. L'utilisateur n'a pas "
    "à s'en préoccuper : pour lui, c'est une seule application."
)
fichier_ref("apps/engine/probare_engine/  — tout le code Python du moteur")
fichier_ref("apps/engine/probare_engine/main.py  — point d'entrée FastAPI")

heading2("2.3 Communication entre les deux parties")
body(
    "L'interface React (dans le navigateur Electron) envoie des requêtes HTTP à l'adresse "
    "locale 127.0.0.1:8767 (visible uniquement sur l'ordinateur, pas sur internet). "
    "Le moteur Python répond avec des données JSON. C'est exactement comme une application web "
    "classique, sauf que tout tourne en local sur la machine de l'auditeur."
)
note_box(
    "🔒 Sécurité : Toutes les données restent sur l'ordinateur. Il n'y a pas de serveur central, "
    "pas de cloud, pas de transmission à un tiers — sauf si l'auditeur active explicitement "
    "l'assistance IA (avec consentement documenté)."
)

heading2("2.4 La base de données")
body(
    "Chaque mission d'audit est stockée dans un fichier SQLite indépendant "
    "(un fichier .db par projet), rangé dans le dossier utilisateur de l'auditeur. "
    "SQLite est une base de données légère qui ne nécessite aucun serveur — c'est simplement "
    "un fichier sur le disque."
)
fichier_ref("apps/engine/probare_engine/storage/db.py  — gestion de la base de données")

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 3. LES 7 ÉTAPES DU PIPELINE
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. Les 7 étapes du pipeline d'audit")

body(
    "Une mission dans Probare suit un cheminement linéaire et irréversible en avant : "
    "on ne peut pas sauter une étape, et on ne peut pas revenir en arrière. "
    "Ce principe garantit la cohérence du dossier d'audit."
)

fichier_ref("apps/engine/probare_engine/statemachine/pipeline.py  — machine à états")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(
    "cadrage  →  ingestion  →  extraction  →  contrôles  →  revue  →  génération  →  opinion"
)
r.font.size = Pt(12)
r.font.bold = True
r.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
p.paragraph_format.space_before = Pt(10)
p.paragraph_format.space_after = Pt(10)

etapes = [
    (
        "1. CADRAGE",
        "L'auditeur crée la mission et renseigne les informations fondamentales : nom du client, "
        "exercice audité (ex : 2025), numéro d'identification (NIF), seuil de signification "
        "(le montant en dessous duquel une erreur est considérée comme négligeable) et seuil "
        "de planification (généralement 60-75% du seuil de signification, utilisé pour les contrôles).",
        "À cette étape, l'auditeur doit également recueillir et documenter le consentement du client "
        "si l'assistance IA doit être utilisée. Sans ce consentement, aucune donnée ne sera envoyée "
        "à un service externe.",
        "apps/desktop/src/renderer/src/pages/Cadrage.tsx"
    ),
    (
        "2. INGESTION",
        "L'auditeur importe les fichiers comptables fournis par le client : la balance générale "
        "(résumé des soldes de tous les comptes) et le grand livre (détail de toutes les écritures "
        "comptables). Les formats acceptés sont Excel (.xlsx, .xls) et CSV.",
        "Chaque cellule importée est immédiatement « étiquetée » avec son origine exacte : "
        "nom du fichier, feuille, numéro de ligne, nom de colonne. Cette étiquette s'appelle "
        "la provenance — elle ne sera jamais perdue.",
        "apps/engine/probare_engine/ingestion/excel_csv.py  +  apps/desktop/src/renderer/src/pages/Ingestion.tsx"
    ),
    (
        "3. EXTRACTION",
        "Phase intermédiaire automatique : le système organise les données importées "
        "pour les rendre utilisables par les contrôles. Cette étape est transparente "
        "pour l'utilisateur.",
        "Les données sont normalisées (virgules → points pour les montants, gestion des espaces, etc.) "
        "et stockées dans la table donnee_sourcee de la base de données.",
        "apps/engine/probare_engine/ingestion/excel_csv.py"
    ),
    (
        "4. CONTRÔLES",
        "C'est le cœur de l'audit. Probare exécute automatiquement jusqu'à 6 contrôles "
        "déterministes sur le cycle trésorerie (et peut en exécuter d'autres sur les cycles "
        "achats et ventes). Chaque contrôle est un calcul Python pur — aucun LLM, "
        "aucune approximation.",
        "Chaque contrôle peut donner un résultat « OK » (pas d'anomalie détectée) ou "
        "déclencher une « exception » (anomalie à examiner). Les exceptions sont classées "
        "par sévérité : mineure, significative ou critique.",
        "apps/engine/probare_engine/controls/engine.py  +  apps/engine/probare_engine/controls/registry.py"
    ),
    (
        "5. REVUE",
        "L'auditeur examine les exceptions levées par les contrôles. Pour chaque exception, "
        "il peut (optionnellement) demander une interprétation par l'IA (Claude) pour obtenir "
        "une explication en langage clair de l'anomalie.",
        "Mais surtout, l'auditeur doit prendre une décision sur chaque exception : "
        "l'accepter (avec justification documentée), la corriger, ou l'escalader. "
        "Cette décision est appelée « tranchement ».",
        "apps/desktop/src/renderer/src/pages/Exceptions.tsx"
    ),
    (
        "6. GÉNÉRATION",
        "Une fois toutes les exceptions tranchées (et seulement à ce moment), l'auditeur "
        "peut générer les livrables. Le logiciel bloque explicitement cette étape tant qu'il "
        "reste des exceptions ouvertes.",
        "Les livrables produits sont : un dossier de travail (.docx) conforme NEP 230, "
        "et un tableau des exceptions (.xlsx). Chaque chiffre du dossier est lié à sa "
        "source — si un chiffre sans provenance est détecté, la génération échoue avec "
        "une erreur explicite.",
        "apps/engine/probare_engine/reporting/export.py  +  apps/desktop/src/renderer/src/pages/Rapport.tsx"
    ),
    (
        "7. OPINION",
        "Étape finale et symbolique : l'auditeur a téléchargé ses livrables. Le dossier "
        "est considéré comme complet. L'opinion d'audit finale (certification, refus de "
        "certifier, certification avec réserves) est rédigée et signée exclusivement "
        "par l'auditeur habilité — Probare ne rédige pas l'opinion.",
        "Cette étape marque la clôture de la mission dans le système.",
        "apps/desktop/src/renderer/src/pages/Rapport.tsx"
    ),
]

for titre, desc1, desc2, fichiers in etapes:
    heading3(titre)
    body(desc1)
    body(desc2)
    fichier_ref(fichiers)
    spacer()

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 4. LE MOTEUR PYTHON
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. Le moteur Python — apps/engine")

body(
    "Le moteur Python est l'ensemble du code qui fait les calculs, gère les données "
    "et produit les livrables. Il est complètement indépendant de l'interface graphique — "
    "on pourrait l'utiliser depuis un terminal sans aucune interface visuelle. "
    "Cette séparation est un choix de conception important : le cœur du logiciel est "
    "testable, prévisible et auditable indépendamment de l'interface."
)


# 4.1 Provenance
heading2("4.1 La provenance des données — DonneeSourcee")
fichier_ref("apps/engine/probare_engine/provenance/models.py")

body(
    "Le concept le plus fondamental de Probare est la DonneeSourcee. C'est une structure "
    "de données immuable (on ne peut pas la modifier après création) qui représente "
    "un seul chiffre ou une seule valeur extraite d'un fichier comptable."
)
body("Chaque DonneeSourcee contient :")
bullet("id — un identifiant unique universel (UUID)")
bullet("projet_id — à quelle mission appartient cette donnée")
bullet("fichier_source_id — de quel fichier elle provient")
bullet("valeur — la valeur elle-même (montant, texte, date...)")
bullet("type — la nature de la valeur (montant, compte, texte, date, numéro de pièce)")
bullet("localisation — l'adresse exacte dans le fichier source (ex: balance_2025:12:Débit)")
bullet("confiance_extraction — niveau de confiance (1.0 = certitude absolue)")
bullet("extrait_par — qui a extrait cette valeur (ingestion-directe pour les fichiers)")
bullet("horodatage — quand elle a été extraite")

body(
    "Principe crucial : aucun nombre ne peut entrer dans un calcul d'audit sans être "
    "une DonneeSourcee. Si lors de l'export un chiffre du rapport n'est pas lié à une "
    "source, le logiciel lève une erreur ProvenanceError et refuse de générer le document.",
    bold_parts=["aucun nombre ne peut entrer dans un calcul d'audit sans être une DonneeSourcee",
                "ProvenanceError"]
)
note_box(
    "💡 Analogie : Imaginez que chaque chiffre dans votre tableur soit accompagné d'un post-it "
    "indiquant exactement dans quel classeur, quel onglet, et quelle case vous l'avez trouvé. "
    "C'est exactement ce que fait DonneeSourcee automatiquement."
)


# 4.2 Ingestion
heading2("4.2 L'ingestion des fichiers comptables")
fichier_ref("apps/engine/probare_engine/ingestion/excel_csv.py")

body(
    "Quand l'auditeur dépose un fichier Excel ou CSV, la fonction lire_fichier() prend en charge "
    "toute la chaîne de traitement :"
)
bullet("Lecture du fichier avec pandas (bibliothèque Python spécialisée dans les tableaux de données)")
bullet("Détection automatique des colonnes : le code recherche des noms de colonnes connus "
       "(compte, débit, crédit, libellé, date, pièce...) avec des synonymes courants")
bullet("Création d'une DonneeSourcee par cellule significative, avec sa localisation précise")
bullet("Calcul du hash SHA-256 du fichier pour détecter toute modification ultérieure")

body(
    "La détection des colonnes est heuristique (basée sur des règles), pas sur l'IA. "
    "Si la détection automatique est incorrecte, l'auditeur peut corriger le mapping manuellement "
    "ou demander une suggestion à l'IA (si le consentement est activé)."
)


# 4.3 Registre des contrôles
heading2("4.3 Le registre des contrôles NEP")
fichier_ref("apps/engine/probare_engine/controls/registry.py")

body(
    "Les contrôles sont déclarés dans un registre — une liste de définitions. "
    "Cette approche est fondamentale : les règles d'audit (références NEP, sévérités, "
    "cycles concernés) sont des données, pas du code en dur. "
    "Ajouter un nouveau contrôle ne nécessite pas de modifier la logique de calcul."
)
body("Les 10 contrôles actuellement enregistrés sont :")

controles = [
    ("TRESOR-BAL-EQUIL", "NEP 500", "Équilibre de la balance",
     "Vérifie que Σ débits = Σ crédits. Si ce n'est pas le cas, la balance est fausse."),
    ("TRESOR-GL-COHER", "NEP 500", "Cohérence grand livre / balance",
     "La somme des mouvements du grand livre par compte doit égaler le solde de la balance."),
    ("TRESOR-ADDITION", "NEP 500", "Contrôle d'addition (foliotage)",
     "Recalcule indépendamment les totaux affichés dans la balance."),
    ("TRESOR-SEQ-PIECES", "NEP 330", "Continuité des séquences de pièces",
     "Détecte les trous (pièce manquante) et doublons dans la numérotation."),
    ("TRESOR-VARIATION", "NEP 520", "Variations N/N-1 significatives",
     "Identifie les variations de solde dépassant le seuil de signification."),
    ("TRESOR-RAPPROCH", "NEP 500", "Rapprochement bancaire",
     "Compare le solde comptable au solde du relevé bancaire."),
    ("ACHAT-SEQ-FACTURES", "NEP 330", "Séquence factures fournisseurs",
     "Trous et doublons dans les numéros de factures fournisseurs."),
    ("ACHAT-VARIATION", "NEP 520", "Variations achats N/N-1",
     "Variations de charges dépassant le seuil de signification."),
    ("VENTE-SEQ-FACTURES", "NEP 330", "Séquence factures clients",
     "Trous et doublons dans les numéros de factures clients."),
    ("VENTE-VARIATION", "NEP 520", "Variations ventes N/N-1",
     "Variations de produits dépassant le seuil de signification."),
]

for ref, nep, libelle, desc in controles:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{ref}  ")
    r1.font.name = "Courier New"
    r1.font.size = Pt(9)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(0x6D, 0x28, 0xD9)
    r2 = p.add_run(f"[{nep}]  ")
    r2.font.size = Pt(9)
    r2.font.bold = True
    r2.font.color.rgb = RGBColor(0x04, 0x78, 0x57)
    r3 = p.add_run(f"{libelle} — {desc}")
    r3.font.size = Pt(10)


# 4.4 Les contrôles déterministes
heading2("4.4 Les contrôles déterministes")
fichier_ref("apps/engine/probare_engine/controls/engine.py")
fichier_ref("apps/engine/tests/test_controls.py  — 20 tests unitaires")

body(
    "Chaque contrôle est une fonction Python pure : pour les mêmes données en entrée, "
    "elle produira toujours exactement le même résultat. Il n'y a aucune aléatoire, "
    "aucune approximation, aucune intervention d'un modèle d'IA dans ces calculs."
)
body("Zoom sur quelques contrôles :")

body("Contrôle d'équilibre (TRESOR-BAL-EQUIL) :", bold_parts=["Contrôle d'équilibre (TRESOR-BAL-EQUIL) :"])
body(
    "Récupère toutes les DonneeSourcee de type « débit » et « crédit » du projet, "
    "fait la somme des deux colonnes, compare. Si l'écart est supérieur à 0,01 €, "
    "une exception est levée avec l'écart exact et les IDs des données sources utilisées."
)

body("Contrôle de séquence des pièces (TRESOR-SEQ-PIECES) :", bold_parts=["Contrôle de séquence des pièces (TRESOR-SEQ-PIECES) :"])
body(
    "Récupère tous les numéros de pièces, les trie numériquement, et cherche "
    "les trous (ex : pièce 1001, 1002, 1004 → la 1003 manque) et les doublons "
    "(même numéro deux fois). Ces anomalies peuvent indiquer une fraude ou une erreur d'enregistrement."
)

body("Contrôle de variations (TRESOR-VARIATION) :", bold_parts=["Contrôle de variations (TRESOR-VARIATION) :"])
body(
    "Compare les soldes des comptes N et N-1 pour identifier les variations "
    "supérieures au seuil de signification défini lors du cadrage. "
    "Une variation de +150% sur un compte de trésorerie mérite une explication."
)

note_box(
    "⚠️ Principe fondamental : si un contrôle produit un résultat incorrect, "
    "c'est un bug de code corrigeable — pas une erreur d'interprétation. "
    "20 tests unitaires vérifient le bon fonctionnement de chaque contrôle "
    "avec des jeux de données équilibrés et déséquilibrés."
)


# 4.5 Machine à états
heading2("4.5 La machine à états du pipeline")
fichier_ref("apps/engine/probare_engine/statemachine/pipeline.py")

body(
    "Une machine à états est un mécanisme qui impose un ordre strict dans lequel "
    "les actions peuvent être effectuées. Chaque mission a un état courant, "
    "et seules certaines transitions sont autorisées."
)
body("Les transitions autorisées sont :")
transitions = [
    ("cadrage", "ingestion", "Mission créée, on peut commencer à importer des fichiers"),
    ("ingestion", "extraction", "Fichiers importés, on normalise les données"),
    ("extraction", "contrôles", "Données prêtes, on lance les contrôles"),
    ("contrôles", "revue", "Contrôles exécutés, l'auditeur examine les exceptions"),
    ("revue", "génération", "Toutes les exceptions tranchées, on génère les livrables"),
    ("génération", "opinion", "Livrables téléchargés, la mission est terminée"),
]
for de, vers, desc in transitions:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{de}  →  {vers}")
    r1.font.name = "Courier New"
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    r2 = p.add_run(f"  ({desc})")
    r2.font.size = Pt(10)

body(
    "La transition vers « génération » a une vérification supplémentaire : "
    "le moteur interroge la base pour s'assurer qu'il ne reste aucune exception "
    "avec le statut « ouverte ». Si c'est le cas, la transition est refusée avec "
    "un message d'erreur explicite.",
    bold_parts=["aucune exception avec le statut « ouverte »"]
)
body(
    "Chaque transition réussie est automatiquement enregistrée dans la piste d'audit "
    "(table journal) avec l'état de départ, l'état d'arrivée, l'acteur, et l'horodatage."
)


# 4.6 Base de données
heading2("4.6 La base de données SQLite")
fichier_ref("apps/engine/probare_engine/storage/db.py")

body(
    "La classe ProjectDB est une couche d'abstraction fine autour de SQLite. "
    "Chaque projet est un fichier .db indépendant, stocké dans "
    "C:/Users/[nom]/.probare/projets/[id-projet]/projet.db"
)
body("La base contient 7 tables :")

tables = [
    ("projet", "Les métadonnées de la mission (nom, client, seuils, état courant)"),
    ("fichier_source", "Les fichiers importés avec leur hash SHA-256"),
    ("donnee_sourcee", "Chaque valeur extraite avec sa provenance complète"),
    ("resultat_calcul", "Les résultats de chaque contrôle (OK ou exception)"),
    ("exception", "Les anomalies détectées, leur sévérité, et les décisions de l'auditeur"),
    ("feuille_travail", "Les feuilles de travail rédigées par l'IA (si utilisée)"),
    ("journal", "La piste d'audit : chaque action horodatée"),
]
for table, desc in tables:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{table}")
    r1.font.name = "Courier New"
    r1.font.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = RGBColor(0x0F, 0x76, 0x6E)
    r2 = p.add_run(f"  —  {desc}")
    r2.font.size = Pt(10)


# 4.7 Anonymisation
heading2("4.7 L'anonymisation pour les appels IA")
fichier_ref("apps/engine/probare_engine/anonymization/anonymizer.py")

body(
    "Quand l'auditeur demande une assistance IA (interprétation d'exception, "
    "rédaction de feuille de travail), les données ne doivent pas révéler l'identité "
    "du client audité. L'anonymiseur résout ce problème."
)
body("Son fonctionnement :")
bullet("Il maintient un mapping local token ↔ valeur réelle (ex : CLIENT-001 ↔ « Société Dupont SARL »)")
bullet("Avant l'envoi à l'API, les identifiants nominatifs sont remplacés par leurs tokens")
bullet("Les montants numériques ne sont PAS anonymisés (l'IA a besoin des vrais chiffres pour interpréter)")
bullet("La correspondance n'est jamais envoyée à l'extérieur — elle reste en mémoire locale")
bullet("Après la réponse de l'IA, les tokens peuvent être re-identifiés si nécessaire")

note_box(
    "🔒 Exemple : au lieu d'envoyer « La société Dupont SARL a un écart de 8 610 000 XDJ », "
    "Probare envoie « CLIENT-001 a un écart de 8 610 000 ». L'IA répond sur CLIENT-001, "
    "et le logiciel affiche la réponse avec le vrai nom."
)


# 4.8 Interface LLM
heading2("4.8 L'interface LLM (Claude)")
fichier_ref("apps/engine/probare_engine/llm/client.py  — interface abstraite")
fichier_ref("apps/engine/probare_engine/llm/claude.py  — implémentation Claude")

body(
    "L'accès à l'IA est structuré autour d'une interface abstraite LLMClient. "
    "Cela signifie que le reste du code ne sait pas qu'il utilise Claude spécifiquement — "
    "il appelle des méthodes standardisées. Demain, on pourrait remplacer Claude par "
    "un autre modèle sans toucher au reste du code."
)
body("L'implémentation Claude (ClaudeClient) propose trois capacités :")
bullet("mapper_colonnes() — utilise claude-haiku-4-5 (modèle léger) pour suggérer "
       "comment faire correspondre les colonnes d'un fichier aux champs attendus")
bullet("interpreter_exception() — utilise claude-sonnet-4-6 (modèle standard) pour "
       "expliquer en langage clair pourquoi une anomalie est suspecte")
bullet("rediger_feuille_travail() — utilise claude-sonnet-4-6 pour rédiger la feuille "
       "de travail du cycle à partir des résultats calculés")

body(
    "Règle absolue : la clé API (ANTHROPIC_API_KEY) est lue uniquement depuis "
    "les variables d'environnement de l'ordinateur. Elle n'est jamais dans le code, "
    "jamais dans la base de données, jamais dans un fichier de configuration.",
    bold_parts=["ANTHROPIC_API_KEY", "jamais dans le code, jamais dans la base de données, jamais dans un fichier de configuration"]
)


# 4.9 Exports
heading2("4.9 Les exports de livrables")
fichier_ref("apps/engine/probare_engine/reporting/export.py")

body(
    "Deux types de livrables sont générés :"
)
bullet("Dossier de travail (.docx) : document Word structuré contenant le récapitulatif "
       "des contrôles, les exceptions et leurs décisions, les feuilles de travail. "
       "Chaque chiffre est accompagné de sa provenance. Si un chiffre sans source est détecté, "
       "le moteur lève une ProvenanceError et refuse de générer le fichier.")
bullet("Tableau des exceptions (.xlsx) : fichier Excel listant toutes les exceptions avec "
       "leur sévérité, description, décision de l'auditeur et horodatage.")

body(
    "Les fichiers sont écrits dans le dossier de données local du projet "
    "(C:/Users/[nom]/.probare/projets/[id]/exports/) et retournés au navigateur "
    "via une réponse HTTP de type « FileResponse »."
)


# 4.10 API REST
heading2("4.10 L'API REST — FastAPI")
fichier_ref("apps/engine/probare_engine/api/routes.py  — toutes les routes")
fichier_ref("apps/engine/probare_engine/main.py  — configuration FastAPI")

body(
    "FastAPI expose un ensemble d'URLs (appelées « routes » ou « endpoints ») "
    "que l'interface React appelle pour toutes les opérations. "
    "Les principales routes sont :"
)

routes = [
    ("GET /api/health", "Vérifie que le moteur Python est bien démarré"),
    ("GET/POST /api/projets", "Liste les projets ou en crée un nouveau"),
    ("GET /api/projets/{id}", "Récupère les détails d'un projet"),
    ("POST /api/projets/{id}/fichiers", "Upload d'un fichier comptable"),
    ("POST /api/projets/{id}/controles/tresorerie", "Lance les 6 contrôles du cycle trésorerie"),
    ("POST /api/projets/{id}/transition", "Avance le pipeline à l'étape suivante"),
    ("POST /api/projets/{id}/exceptions/{eid}/trancher", "Documente la décision sur une exception"),
    ("POST /api/projets/{id}/exceptions/{eid}/interpreter", "Demande une interprétation IA"),
    ("POST /api/projets/{id}/generer-feuille", "Rédige une feuille de travail avec l'IA"),
    ("POST /api/projets/{id}/exporter-dossier", "Génère et télécharge le .docx"),
    ("POST /api/projets/{id}/exporter-exceptions", "Génère et télécharge le .xlsx"),
    ("GET /api/projets/{id}/journal", "Récupère la piste d'audit complète"),
]
for route, desc in routes:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f"{route}")
    r1.font.name = "Courier New"
    r1.font.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(0x04, 0x78, 0x57)
    r2 = p.add_run(f"  —  {desc}")
    r2.font.size = Pt(10)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 5. L'INTERFACE DESKTOP
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. L'interface desktop — apps/desktop")

body(
    "L'interface desktop est ce que l'auditeur voit et manipule. Elle est "
    "construite comme une application web moderne (React + TypeScript) "
    "embarquée dans une fenêtre native grâce à Electron."
)


# 5.1 Processus principal
heading2("5.1 Le processus principal Electron")
fichier_ref("apps/desktop/src/main/index.ts")

body(
    "Electron fonctionne avec deux processus séparés : le « main process » (processus principal) "
    "qui contrôle la fenêtre, le système de fichiers et lance le sidecar Python ; "
    "et le « renderer process » (processus de rendu) qui affiche l'interface web."
)
body(
    "Le processus principal crée la fenêtre de l'application, démarre le sidecar Python, "
    "et gère l'arrêt propre (en tuant le sidecar quand l'application se ferme)."
)


# 5.2 Sidecar
heading2("5.2 Le sidecar Python")
fichier_ref("apps/desktop/src/main/sidecar.ts")

body(
    "Le sidecar est le mécanisme par lequel l'application de bureau lance et gère "
    "le moteur Python. Voici comment il fonctionne au démarrage :"
)
bullet("1. Recherche un port libre (à partir de 8765) pour éviter les conflits")
bullet("2. Lance python -m uvicorn probare_engine.main:app --port [port trouvé]")
bullet("3. Interroge toutes les 300ms l'URL /api/health jusqu'à obtenir une réponse 200")
bullet("4. Une fois le sidecar prêt, notifie l'interface React du port à utiliser")
bullet("5. À la fermeture, envoie SIGTERM puis SIGKILL (après 3 secondes) au processus")

body(
    "En mode développement, le code Python est lancé directement depuis le dossier source. "
    "En production (application installée), un exécutable PyInstaller embarqué dans "
    "les ressources de l'application est utilisé."
)


# 5.3 Interface React
heading2("5.3 L'interface React")
fichier_ref("apps/desktop/src/renderer/src/App.tsx  — routeur principal")
fichier_ref("apps/desktop/src/renderer/src/stores/projetStore.ts  — état global")
fichier_ref("apps/desktop/src/renderer/src/hooks/useApi.ts  — appels HTTP")

body(
    "L'interface est organisée autour de trois concepts :"
)
bullet("Le routeur (React Router v6) : associe chaque URL à une page. "
       "ex: /#/projet/abc123/controles → affiche la page Contrôles pour le projet abc123")
bullet("Le store global (Zustand) : stocke l'état partagé entre les pages "
       "(projet actif, liste de fichiers, résultats de contrôles, exceptions). "
       "Quand une donnée change, tous les composants qui l'utilisent se mettent à jour automatiquement.")
bullet("Le hook useApi : centralise tous les appels HTTP vers le moteur Python. "
       "Toutes les pages passent par ce hook pour GET, POST, PATCH et téléchargements.")

body(
    "Un hook spécial useSyncProjet() est appelé au montage de chaque page projet. "
    "Il recharge l'état du projet depuis l'API pour s'assurer que l'état affiché "
    "est toujours à jour, même si une transition a eu lieu dans un autre onglet.",
    bold_parts=["useSyncProjet()"]
)
fichier_ref("apps/desktop/src/renderer/src/hooks/useProjet.ts  — synchronisation état")


# 5.4 Pages une par une
heading2("5.4 Les pages de l'application")

pages = [
    (
        "Dashboard",
        "apps/desktop/src/renderer/src/pages/Dashboard.tsx",
        "Page d'accueil. Affiche la liste de toutes les missions ouvertes. "
        "Chaque mission est présentée comme une carte avec son nom, son client, "
        "son exercice et son état courant dans le pipeline. "
        "Un bouton « Nouvelle mission » permet de créer un projet et de passer à la page Cadrage."
    ),
    (
        "Cadrage",
        "apps/desktop/src/renderer/src/pages/Cadrage.tsx",
        "Formulaire de paramétrage de la mission. L'auditeur renseigne le nom du client, "
        "le NIF, l'exercice, les seuils de signification et de planification. "
        "Un switch de consentement (avec horodatage automatique) permet d'activer "
        "les fonctionnalités IA pour cette mission."
    ),
    (
        "Ingestion",
        "apps/desktop/src/renderer/src/pages/Ingestion.tsx",
        "Zone de dépôt de fichiers (drag & drop ou parcourir). L'auditeur importe "
        "séparément la balance et le grand livre. Après chaque import, "
        "un résumé de l'extraction (nombre de données extraites, colonnes détectées) "
        "est affiché. Une modale de vérification du mapping permet de corriger "
        "les associations colonnes/champs si nécessaire."
    ),
    (
        "Contrôles",
        "apps/desktop/src/renderer/src/pages/Controles.tsx",
        "Tableau de bord des contrôles. Un bouton « Lancer les contrôles » "
        "déclenche l'exécution de tous les contrôles du cycle trésorerie. "
        "Les résultats apparaissent en liste : une icône verte pour « OK », "
        "rouge pour « Exception ». Chaque résultat est cliquable pour voir "
        "les détails, la valeur calculée et les IDs des données sources utilisées."
    ),
    (
        "Exceptions",
        "apps/desktop/src/renderer/src/pages/Exceptions.tsx",
        "Revue des anomalies. Chaque exception est affichée avec sa sévérité "
        "(code couleur : rouge critique, orange significatif, gris mineur), "
        "sa description et la norme NEP concernée. "
        "Pour chaque exception ouverte, deux actions sont disponibles : "
        "« Interpréter avec l'IA » (requiert consentement) et « Trancher cette exception » "
        "(ouvre une modale pour saisir la décision et le nom de l'auditeur). "
        "Quand toutes les exceptions sont tranchées, un bandeau vert apparaît "
        "et le bouton « Générer le dossier » devient actif."
    ),
    (
        "Rapport",
        "apps/desktop/src/renderer/src/pages/Rapport.tsx",
        "Page de génération des livrables. Affiche le récapitulatif "
        "(nombre de contrôles, résultats OK, exceptions). "
        "Deux boutons d'export : dossier de travail (.docx) et tableau des exceptions (.xlsx). "
        "Optionnellement, un bouton « Rédiger avec l'IA » permet de générer "
        "une feuille de travail narrative (requiert ANTHROPIC_API_KEY et consentement). "
        "Le téléchargement du dossier fait automatiquement transitionner la mission "
        "vers l'état « opinion »."
    ),
    (
        "Journal",
        "apps/desktop/src/renderer/src/pages/Journal.tsx",
        "Piste d'audit complète. Liste chronologique (inverse) de toutes les actions : "
        "transitions d'état, tranchements d'exceptions, exports, appels IA. "
        "Chaque entrée montre l'horodatage, le type d'événement et les données associées "
        "en JSON. C'est la preuve irréfutable de ce qui s'est passé dans la mission."
    ),
]

for nom, fichier, desc in pages:
    heading3(nom)
    fichier_ref(fichier)
    body(desc)
    spacer()

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 6. FLUX COMPLET D'UNE MISSION
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. Le flux complet d'une mission d'audit — exemple concret")

body(
    "Voici le cheminement complet d'une mission fictive pour mieux comprendre "
    "comment toutes les parties s'articulent."
)

etapes_flux = [
    (
        "Lundi matin — L'auditeur ouvre Probare",
        "Electron démarre. Le main process (index.ts) cherche un port libre "
        "et lance python -m uvicorn sur le port 8767. Le sidecar (sidecar.ts) "
        "interroge /api/health toutes les 300ms. Après ~2 secondes, la réponse "
        "{\"status\": \"ok\"} arrive. L'interface React s'affiche — l'auditeur voit "
        "le tableau de bord avec ses missions."
    ),
    (
        "Création de la mission",
        "L'auditeur clique « Nouvelle mission ». Dashboard.tsx envoie POST /api/projets "
        "avec le nom de la mission. Le moteur crée un UUID, initialise une base SQLite "
        "dans ~/.probare/projets/[uuid]/projet.db, enregistre l'état « cadrage » "
        "et renvoie le projet créé. L'interface navigue vers /#/projet/[uuid]/cadrage."
    ),
    (
        "Cadrage — paramétrage",
        "Sur la page Cadrage.tsx, l'auditeur remplit : client « Société Exemple SA », "
        "exercice 2025, seuil de signification 500 000 XDJ. Il active le consentement. "
        "PATCH /api/projets/[id] enregistre ces données. La machine à états "
        "reste en « cadrage » — l'auditeur passe à l'étape suivante manuellement "
        "en cliquant « Commencer l'ingestion »."
    ),
    (
        "Ingestion — import des fichiers",
        "L'auditeur dépose balance_2025.xlsx. Ingestion.tsx envoie un FormData "
        "en POST /api/projets/[id]/fichiers. Le moteur appelle lire_fichier() "
        "depuis excel_csv.py : pandas lit le fichier, détecte automatiquement "
        "les colonnes (compte → « N° Cpte », débit → « Mouvements D »...), "
        "crée 1 240 DonneeSourcee avec localisation exacte (ex: balance_2025!Feuil1:5:Mouvements D), "
        "et retourne le résumé. La même opération est répétée pour grand_livre_2025.xlsx."
    ),
    (
        "Lancement des contrôles",
        "L'auditeur clique « Lancer les contrôles ». POST /api/projets/[id]/controles/tresorerie "
        "est appelé. Le moteur exécute successivement les 6 fonctions de controls/engine.py. "
        "Résultat : 5 contrôles OK, 1 exception (TRESOR-GL-COHER — écart de 8 610 000 XDJ "
        "entre le grand livre et la balance sur le compte 512). "
        "L'exception est enregistrée dans la table exception avec sévérité « critique »."
    ),
    (
        "Revue — tranchement de l'exception",
        "Sur la page Exceptions.tsx, l'auditeur voit l'exception rouge. "
        "Il clique « Trancher cette exception », saisit sa décision : "
        "« Écart vérifié — écriture de régularisation de clôture n°2025/12/089 "
        "passée après l'arrêté de la balance provisoire. Pièce reçue le 05/01/2026. » "
        "POST /api/projets/[id]/exceptions/[eid]/trancher met à jour le statut "
        "en « tranchee » et log l'action dans le journal. "
        "Le bandeau vert « Toutes les exceptions ont été tranchées » apparaît."
    ),
    (
        "Génération des livrables",
        "L'auditeur clique « Générer le dossier ». Deux transitions sont effectuées "
        "via POST /api/projets/[id]/transition : controles → revue, puis revue → génération. "
        "La seconde transition vérifie qu'il n'y a plus d'exceptions ouvertes. "
        "L'interface navigue vers la page Rapport. L'auditeur clique "
        "« Dossier de travail » — POST /api/projets/[id]/exporter-dossier "
        "appelle generer_dossier_travail() depuis export.py, écrit le .docx, "
        "et le retourne en FileResponse. Le navigateur déclenche le téléchargement."
    ),
    (
        "Clôture",
        "Le téléchargement du dossier déclenche la transition génération → opinion. "
        "La mission passe à son état final. L'auditeur peut consulter "
        "le Journal pour voir la trace complète de toutes les actions."
    ),
]

for titre, desc in etapes_flux:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(titre)
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x3B, 0x30, 0xA8)
    body(desc)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 7. RÈGLES NON NÉGOCIABLES
# ══════════════════════════════════════════════════════════════════════════════
heading1("7. Les règles non négociables (CLAUDE.md)")
fichier_ref("probare/CLAUDE.md  — règles permanentes du projet")

body(
    "Ces règles ont été définies lors de la conception et s'appliquent à tous les "
    "développements présents et futurs du logiciel, sans exception."
)

regles = [
    (
        "Règle 1 — Le LLM ne calcule jamais",
        "Toute arithmétique (additions, comparaisons, pourcentages, seuils) "
        "est implémentée en code Python. L'IA est convoquée uniquement pour "
        "du texte : expliquer, résumer, rédiger. Jamais pour un chiffre. "
        "Raison : un LLM peut halluciner un montant. Un bug Python s'identifie et se corrige."
    ),
    (
        "Règle 2 — Pas de valeur sans provenance",
        "Tout nombre dans un calcul ou un livrable doit être une DonneeSourcee "
        "avec un lien vers son fichier d'origine. ProvenanceError bloque la génération "
        "si ce n'est pas respecté. Raison : un auditeur doit pouvoir justifier chaque "
        "chiffre devant un tribunal."
    ),
    (
        "Règle 3 — Clé API en variable d'environnement uniquement",
        "ANTHROPIC_API_KEY n'est jamais dans le code, jamais dans la base, "
        "jamais dans un fichier de config versionné. Raison : une clé committée "
        "accidentellement dans Git expose le compte et peut générer des coûts "
        "non maîtrisés."
    ),
    (
        "Règle 4 — LLM derrière une interface",
        "L'implémentation Claude est derrière LLMClient. Le reste du code "
        "ne sait pas qu'il utilise Claude. Raison : permet de changer de fournisseur "
        "LLM sans refactoring, et de mocker l'IA dans les tests."
    ),
    (
        "Règle 5 — NEP en données, pas en dur",
        "Les contrôles et leurs références NEP sont déclarés dans le registre (registry.py), "
        "pas codés en dur dans la logique métier. Raison : les normes évoluent — "
        "un nouveau contrôle s'ajoute en quelques lignes sans toucher au moteur."
    ),
    (
        "Règle 6 — Tout est journalisé",
        "Chaque transition d'état et chaque appel LLM est écrit dans la table journal "
        "avec horodatage. Raison : conformité NEP 230 — le dossier doit témoigner de "
        "tout ce qui s'est passé."
    ),
    (
        "Règle 7 — Tests pour chaque contrôle déterministe",
        "Chaque fonction de contrôle a au minimum un test « cas OK » "
        "et un test « cas exception ». Raison : les contrôles sont la raison d'être "
        "du logiciel — un bug ici est une faute professionnelle."
    ),
    (
        "Règle 8 — UI et livrables en français",
        "Tous les textes affichés, tous les libellés, tous les documents générés "
        "sont en français. Raison : le logiciel cible des auditeurs francophones "
        "et les normes NEP sont en français."
    ),
]

for titre, desc in regles:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run(titre)
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    body(desc)

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# 8. GLOSSAIRE
# ══════════════════════════════════════════════════════════════════════════════
heading1("8. Glossaire")

termes = [
    ("API REST", "Interface de communication entre deux logiciels via HTTP. Probare utilise FastAPI pour exposer ses fonctionnalités au frontend."),
    ("Balance générale", "Document comptable résumant le solde débiteur et créditeur de chaque compte à une date donnée."),
    ("DonneeSourcee", "Structure de données de Probare représentant une valeur extraite avec sa provenance complète (fichier, ligne, colonne)."),
    ("Electron", "Technologie permettant de créer des applications de bureau avec des technologies web (HTML/CSS/JS)."),
    ("Exception", "Anomalie détectée par un contrôle déterministe, qui doit être examinée et tranchée par l'auditeur."),
    ("FastAPI", "Framework Python pour créer des API REST rapidement et avec typage fort."),
    ("Grand livre", "Document comptable listant le détail de toutes les écritures comptables par compte."),
    ("LLMClient", "Interface abstraite de Probare pour les appels aux modèles de langage (IA)."),
    ("Machine à états", "Mécanisme qui impose un ordre strict dans les transitions entre les étapes d'une mission."),
    ("Monorepo", "Structure de dépôt Git contenant plusieurs applications indépendantes (ici : desktop + engine)."),
    ("NEP", "Norme d'Exercice Professionnel — règles que les auditeurs légaux français doivent respecter."),
    ("Provenance", "Lien traçable entre une valeur et son fichier source (fichier, feuille, ligne, colonne)."),
    ("ProvenanceError", "Erreur levée si un livrable tente d'inclure un chiffre sans provenance traçable."),
    ("React", "Bibliothèque JavaScript pour construire des interfaces utilisateur à base de composants."),
    ("Seuil de signification", "Montant en dessous duquel une erreur est considérée comme non significative pour l'opinion d'audit."),
    ("Sidecar", "Processus auxiliaire (ici : le moteur Python) lancé et géré par l'application principale (Electron)."),
    ("SQLite", "Base de données légère stockée dans un fichier, sans serveur nécessaire."),
    ("Tranchement", "Décision documentée de l'auditeur sur une exception : acceptée, corrigée, ou escaladée."),
    ("TypeScript", "JavaScript avec un système de types fort, pour détecter les erreurs à la compilation."),
    ("Zustand", "Bibliothèque légère de gestion d'état global pour les applications React."),
]

for terme, definition in termes:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{terme}  —  ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    r2 = p.add_run(definition)
    r2.font.size = Pt(11)


# ── Pied de page ───────────────────────────────────────────────────────────────
spacer()
spacer()
separator()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Probare v0.1 — Documentation MVP — Confidentiel — Juin 2026")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)


# ── Sauvegarde ─────────────────────────────────────────────────────────────────
output_path = r"D:\projet\Audit_Comptable\Documentation_Probare_MVP.docx"
doc.save(output_path)
print(f"Document généré : {output_path}")
