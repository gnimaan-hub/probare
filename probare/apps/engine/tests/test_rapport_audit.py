"""Tests de la phase Rapport d'audit : opinion (CRUD), référentiel comptable,
synthèse de mission et génération du rapport d'audit + mémorandum."""
from __future__ import annotations
import uuid
import pytest

from probare_engine.storage.db import ProjectDB
from probare_engine.normes import libelle_referentiel_comptable, REFERENTIELS_COMPTABLES
from probare_engine.reporting.export import (
    generer_rapport_audit, generer_memorandum_controle_comptes,
)


@pytest.fixture
def db(tmp_path):
    d = ProjectDB(tmp_path / "audit.db")
    d.connect()
    yield d
    d.close()


def _projet(db, **kw):
    pid = str(uuid.uuid4())
    db.create_projet({"id": pid, "nom": "HARBI MATERIAUX SARL",
                      "client": "HARBI MATERIAUX SARL", "exercice": "2025",
                      "cycles_couverts": ["tresorerie", "ventes"], **kw})
    return pid


# ─── Référentiel comptable ───────────────────────────────────────────────────

def test_referentiel_comptable_defaut(db):
    pid = _projet(db)
    assert db.get_projet(pid)["referentiel_comptable"] == "pcgd"


def test_referentiel_comptable_persiste(db):
    pid = _projet(db, referentiel_comptable="ifrs")
    assert db.get_projet(pid)["referentiel_comptable"] == "ifrs"
    updated = db.update_projet(pid, {"referentiel_comptable": "syscohada"})
    assert updated["referentiel_comptable"] == "syscohada"


def test_libelle_referentiel_comptable():
    assert "Djibouti" in libelle_referentiel_comptable("pcgd")
    assert "IFRS" in libelle_referentiel_comptable("ifrs")
    # code inconnu → défaut PCGD
    assert libelle_referentiel_comptable("xxx") == REFERENTIELS_COMPTABLES["pcgd"]
    assert libelle_referentiel_comptable(None) == REFERENTIELS_COMPTABLES["pcgd"]


# ─── Opinion : CRUD upsert ───────────────────────────────────────────────────

def test_opinion_absente_par_defaut(db):
    pid = _projet(db)
    assert db.get_opinion(pid) is None


def test_opinion_upsert_merge_preserve_champs(db):
    pid = _projet(db)
    op = db.save_opinion(pid, {
        "rigueur": "moderee", "type_opinion": "avec_reserve",
        "titre": "Opinion avec réserve", "texte_opinion": "À notre avis, sous réserve...",
        "fondement": "Selon les normes ISA.", "justification": "Cumul > seuil.",
        "proposee_par_ia": 1, "modele_ia": "claude-opus-4-8", "validee": 0,
    })
    assert op["type_opinion"] == "avec_reserve"
    assert op["validee"] == 0
    genere_le = op["genere_le"]

    # Mise à jour partielle : ne doit pas écraser les champs non fournis
    op2 = db.save_opinion(pid, {"validee": 1, "validee_par": "Auditeur"})
    assert op2["validee"] == 1
    assert op2["validee_par"] == "Auditeur"
    assert op2["texte_opinion"] == "À notre avis, sous réserve..."
    assert op2["type_opinion"] == "avec_reserve"
    assert op2["genere_le"] == genere_le  # date de création préservée


# ─── Synthèse de mission (récapitulatif déterministe) ────────────────────────

def test_synthese_mission_structure(db):
    from probare_engine.api.routes import _synthese_mission
    pid = _projet(db, seuil_signification=500_000, seuil_planification=375_000)
    db.save_resultat({"id": str(uuid.uuid4()), "projet_id": pid,
                      "controle_ref": "TRESOR-BAL-EQUIL", "valeur": 0.0,
                      "statut": "ok", "details": "OK", "sources": ["balance.csv"]})
    projet = db.get_projet(pid)
    recap = _synthese_mission(db, pid, projet)
    assert recap["projet"]["referentiel_comptable"] == "pcgd"
    assert "Djibouti" in recap["projet"]["referentiel_comptable_label"]
    ids_phases = {p["id"] for p in recap["phases"]}
    assert {"cadrage", "travaux_substantifs", "revue", "opinion"} <= ids_phases
    assert "ingérés" in recap["fichiers"] and "produits" in recap["fichiers"]
    assert recap["anomalies"]["cumul_non_corrigees"] == 0.0


# ─── Génération des livrables finaux ─────────────────────────────────────────

def _docx_text(path):
    """Texte de tout le document, y compris les cellules de tableaux (la nouvelle
    charte rend page de garde, bandeaux de section et encadrés via des tableaux)."""
    import docx

    def _cell_text(cell):
        parts = [p.text for p in cell.paragraphs]
        for t in cell.tables:
            for row in t.rows:
                for c in row.cells:
                    parts.append(_cell_text(c))
        return "\n".join(parts)

    doc = docx.Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                parts.append(_cell_text(c))
    return "\n".join(parts)


CABINET = {
    "nom": "Cabinet NIMAAN & Associés", "forme_juridique": "SARL",
    "adresse_ville": "Djibouti", "adresse_pays": "Djibouti",
    "numero_ordre": "CAC-DJ-042", "responsable_nom": "Gouled Ahmed",
    "responsable_titre": "Commissaire aux comptes",
}


def test_generer_rapport_audit(db, tmp_path):
    pid = _projet(db, referentiel_comptable="pcgd")
    db.update_planification(pid, {"forme_juridique": "SARL",
                                  "dirigeants": [{"nom": "M. X", "fonction": "Gérant"}]})
    projet = db.get_projet(pid)
    opinion = {"type_opinion": "sans_reserve", "titre": "Opinion sans réserve",
               "texte_opinion": "À notre avis, les comptes présentent sincèrement...",
               "fondement": "Selon les normes ISA.", "observations": "", "validee": 1}
    out = generer_rapport_audit(projet, opinion, tmp_path / "rapport.docx",
                                cabinet=CABINET, plan=db.get_or_create_planification(pid))
    assert out.exists() and out.stat().st_size > 5000
    txt = _docx_text(out)
    # Bandeaux de section rendus en capitales par la charte ; le fond (opinion,
    # destinataire, référentiel, signature) reste en casse naturelle.
    for needle in ["Cabinet NIMAAN", "OPINION SANS RÉSERVE", "FONDEMENT DE L'OPINION",
                   "Plan Comptable Général de Djibouti", "M. X, Gérant",
                   "Il appartient à la direction", "Commissaire aux comptes"]:
        assert needle in txt, f"absent du rapport : {needle}"


def test_generer_memorandum_triptyque(db, tmp_path):
    pid = _projet(db)
    db.save_resultat({"id": str(uuid.uuid4()), "projet_id": pid,
                      "controle_ref": "VENTE-CUT-OFF", "valeur": 0.42,
                      "statut": "exception", "details": "42% en fin d'exercice",
                      "sources": ["gl.csv"]})
    eid = str(uuid.uuid4())
    db.save_exception({"id": eid, "projet_id": pid, "controle_ref": "VENTE-CUT-OFF",
                       "nep_ref": "NEP 330", "severite": "significative",
                       "description": "Concentration ventes décembre", "statut": "ouverte"})
    db.trancher_exception(eid, "Non corrigée", "Auditeur",
                          type_resolution="non_corrigee", montant_incidence=120_000)
    db.save_feuille_travail({"id": str(uuid.uuid4()), "projet_id": pid, "cycle": "ventes",
                             "contenu_redige": "Conclusion : anomalie non corrigée.",
                             "sources": ["VENTE-CUT-OFF"], "nep_ref": "NEP 330"})
    projet = db.get_projet(pid)
    from probare_engine.api.routes import _synthese_mission
    synthese = _synthese_mission(db, pid, projet)
    out = generer_memorandum_controle_comptes(
        projet, db.list_resultats(pid), db.list_exceptions(pid), db.list_feuilles(pid),
        tmp_path / "memo.docx", plan=db.get_or_create_planification(pid),
        controles_ignores=[{"controle_ref": "TRESOR-RAPPROCH", "cycle": "tresorerie",
                            "raison": "Relevé bancaire non fourni"}],
        cabinet=CABINET, synthese=synthese)
    assert out.exists() and out.stat().st_size > 5000
    txt = _docx_text(out)
    # Nouvelle charte : mémo COMPLET rédigé en langage naturel (plus de sous-titres
    # rigides « Objectifs / Travaux effectués »), page de garde + déroulé de mission.
    for needle in ["Mémorandum", "Prise de connaissance et cadrage",
                   "Cycle Ventes-Clients", "assertions d'audit", "VENTE-CUT-OFF",
                   "non corrigée", "supports de travail", "TRESOR-RAPPROCH"]:
        assert needle in txt, f"absent du mémorandum : {needle}"


# ─── Régression : feuille de travail écrite en JSON brut dans le mémorandum ──
#
# Deux causes distinctes constatées en production :
# 1) `_parse_json` rejetait tout le document dès qu'une valeur de chaîne
#    contenait un saut de ligne littéral (non échappé en `\n`) — fréquent sur
#    un champ "contenu" très long et très formaté. Le code de repli
#    enregistrait alors le JSON brut (non parsé) comme contenu rédigé.
# 2) Chaque régénération d'une feuille pour un cycle donné s'ADDITIONNAIT aux
#    précédentes au lieu de les remplacer, faisant apparaître des doublons
#    (dont d'anciennes versions cassées) dans le dossier de travail.

def _claude_client():
    import os
    from probare_engine.llm.claude import ClaudeClient
    os.environ.setdefault("ANTHROPIC_API_KEY", "test-key-unused")
    return ClaudeClient()


def test_parse_json_tolere_saut_de_ligne_litteral():
    """Un JSON par ailleurs valide, mais avec un caractère de contrôle brut
    (vrai \\n, non échappé) dans une valeur de chaîne, doit tout de même
    être décodé plutôt que rejeté en bloc."""
    c = _claude_client()
    texte_brut_avec_retour_ligne_litteral = (
        '```json\n{"titre": "Feuille de travail", "contenu": "Ligne 1\n'
        'Ligne 2 avec un retour litteral", "nep_refs": [], "conclusion": "sans_reserve"}\n```'
    )
    result = c._parse_json(texte_brut_avec_retour_ligne_litteral)
    assert isinstance(result, dict), "le JSON aurait dû être décodé malgré le saut de ligne littéral"
    assert result["titre"] == "Feuille de travail"
    assert "Ligne 1" in result["contenu"] and "Ligne 2" in result["contenu"]


def test_parse_json_toujours_none_si_reellement_invalide():
    """Un texte qui n'est pas du JSON (même tolérant) doit rester None —
    pas de faux positif qui masquerait un vrai échec de génération."""
    c = _claude_client()
    assert c._parse_json("Ceci n'est pas du JSON du tout.") is None


def test_regeneration_feuille_remplace_au_lieu_de_saccumuler(db):
    """delete_feuilles_par_cycle purge les anciennes feuilles d'un cycle avant
    l'écriture de la nouvelle — pas d'accumulation de doublons."""
    pid = _projet(db)
    db.save_feuille_travail({"id": str(uuid.uuid4()), "projet_id": pid, "cycle": "tresorerie",
                             "contenu_redige": "```json\n{\"titre\": ...(tronqué)", "sources": [],
                             "nep_ref": "ISA 230"})
    assert len(db.list_feuilles(pid)) == 1

    db.delete_feuilles_par_cycle(pid, "tresorerie")
    db.save_feuille_travail({"id": str(uuid.uuid4()), "projet_id": pid, "cycle": "tresorerie",
                             "contenu_redige": "Feuille de travail correctement rédigée.",
                             "sources": [], "nep_ref": "ISA 230"})

    feuilles = db.list_feuilles(pid)
    assert len(feuilles) == 1, "l'ancienne feuille du cycle aurait dû être remplacée, pas dupliquée"
    assert feuilles[0]["contenu_redige"] == "Feuille de travail correctement rédigée."


def test_delete_feuilles_par_cycle_isole_les_autres_cycles(db):
    """La purge ne doit affecter que le cycle régénéré, pas les autres."""
    pid = _projet(db)
    db.save_feuille_travail({"id": str(uuid.uuid4()), "projet_id": pid, "cycle": "tresorerie",
                             "contenu_redige": "Trésorerie", "sources": [], "nep_ref": "ISA 230"})
    db.save_feuille_travail({"id": str(uuid.uuid4()), "projet_id": pid, "cycle": "ventes",
                             "contenu_redige": "Ventes", "sources": [], "nep_ref": "ISA 230"})

    db.delete_feuilles_par_cycle(pid, "tresorerie")

    feuilles = db.list_feuilles(pid)
    assert len(feuilles) == 1
    assert feuilles[0]["cycle"] == "ventes"
