"""Tests d'intégration des contrôles de détail dans le dossier de travail (.docx).

Requiert python-docx (ignoré sinon). Vérifie que la circularisation (NEP 505)
et les sondages (NEP 530) apparaissent dans le livrable et que la provenance
des soldes comptables est contrôlée.
"""
import pytest

docx = pytest.importorskip("docx")

from probare_engine.reporting.export import generer_dossier_travail, ProvenanceError


PROJET = {
    "client": "Société Test",
    "exercice": "2025",
    "seuil_signification": 1_000_000.0,
}


def _lire_texte(path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            parts.extend(c.text for c in row.cells)
    return "\n".join(parts)


def test_dossier_inclut_circularisation(tmp_path):
    circs = [{
        "compte": "411001", "libelle": "Client Alpha",
        "solde_comptable": 5_000_000.0, "solde_confirme": 5_000_000.0,
        "ecart": 0.0, "ecart_pct": 0.0, "est_significatif": False,
        "statut": "clos", "sources": ["src-1"],
        "analyse_ia": '{"synthese": "Solde confirmé sans écart.", "diligences": []}',
    }]
    out = tmp_path / "dossier.docx"
    generer_dossier_travail(PROJET, [], [], [], out, circularisations=circs)
    texte = _lire_texte(out)
    assert "Circularisation des tiers (NEP 505)" in texte
    assert "411001" in texte
    assert "Client Alpha" in texte


def test_dossier_inclut_sondage(tmp_path):
    sondages = [{
        "libelle": "Sondage achats", "cycle": "achats",
        "population": 500, "taille_echantillon": 50, "niveau_confiance": 95,
        "montant_population": 10_000_000.0, "nb_anomalies": 2,
        "taux_anomalie": 0.04, "montant_projete": 200_000.0,
        "conclusion_ia": '{"synthese": "Taux acceptable.", "diligences": ["Vérifier 2 pièces."], "impact_opinion": "Aucun."}',
    }]
    out = tmp_path / "dossier.docx"
    generer_dossier_travail(PROJET, [], [], [], out, sondages=sondages)
    texte = _lire_texte(out)
    assert "Sondages sur pièces (NEP 530)" in texte
    assert "Sondage achats" in texte
    assert "Taux acceptable." in texte


def test_dossier_sans_detail_pas_de_sections(tmp_path):
    """Sans circularisation ni sondage, les sections 4 et 5 sont absentes."""
    out = tmp_path / "dossier.docx"
    generer_dossier_travail(PROJET, [], [], [], out)
    texte = _lire_texte(out)
    assert "NEP 505" not in texte
    assert "NEP 530" not in texte


def test_circularisation_solde_sans_source_bloque(tmp_path):
    """Un solde comptable circularisé sans provenance bloque la génération (règle B.2)."""
    circs = [{
        "compte": "411001", "libelle": "Client X",
        "solde_comptable": 5_000_000.0, "sources": [],
        "statut": "propose",
    }]
    out = tmp_path / "dossier.docx"
    with pytest.raises(ProvenanceError):
        generer_dossier_travail(PROJET, [], [], [], out, circularisations=circs)


def test_ecart_significatif_marque(tmp_path):
    circs = [{
        "compte": "411002", "libelle": "Client Béta",
        "solde_comptable": 5_000_000.0, "solde_confirme": 4_000_000.0,
        "ecart": 1_000_000.0, "ecart_pct": 20.0, "est_significatif": True,
        "statut": "reponse_recue", "sources": ["src-2"],
    }]
    out = tmp_path / "dossier.docx"
    generer_dossier_travail(PROJET, [], [], [], out, circularisations=circs)
    texte = _lire_texte(out)
    assert "⚠" in texte  # marqueur d'écart significatif
