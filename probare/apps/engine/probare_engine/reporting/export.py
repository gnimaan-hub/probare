"""Export dossier de travail (docx) et tableaux (xlsx) avec contrôle de provenance."""
from __future__ import annotations
import io
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from ..normes import norme, prefixe_actif, libelle_referentiel_comptable


def _now() -> str:
    return datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M UTC")


class ProvenanceError(Exception):
    """Levée si un chiffre sans source est détecté dans un livrable."""


def _verifier_sources(valeur: Any, sources: list[str], libelle: str) -> None:
    """Échoue si une valeur numérique n'a pas de source."""
    if isinstance(valeur, (int, float)) and valeur != 0 and not sources:
        raise ProvenanceError(
            f"Valeur '{libelle}' = {valeur} sans source. "
            f"La génération est bloquée (règle B.2)."
        )


_MONTANT_RE = re.compile(r'\b\d{1,3}(?:[,\s]\d{3})+(?:[.,]\d+)?\b|\b\d{5,}(?:[.,]\d+)?\b')


def _verifier_provenance_texte_llm(
    contenu: str,
    valeurs_python: set[float],
    contexte: str,
) -> list[str]:
    """
    Scanne contenu à la recherche de grands nombres (≥ 1 000) non calculés par Python.
    Retourne la liste des avertissements (n'interrompt pas la génération).
    """
    avertissements = []
    for m in _MONTANT_RE.finditer(contenu):
        raw = m.group().replace(',', '').replace(' ', '').replace('\xa0', '')
        try:
            val = float(raw)
        except ValueError:
            continue
        if val < 1000:
            continue
        # Tolérance ±0.5% ou 1 FDJ
        if not any(abs(val - v) <= max(0.005 * v, 1.0) for v in valeurs_python):
            avertissements.append(
                f"Nombre {val:,.0f} dans '{contexte}' non retrouvé dans les valeurs Python."
            )
    return avertissements


def generer_dossier_travail(
    projet: dict,
    resultats: list[dict],
    exceptions: list[dict],
    feuilles: list[dict],
    output_path: Path,
    controles_ignores: list[dict] | None = None,
    synthese_anomalies: dict | None = None,
) -> Path:
    """Génère le dossier de travail en .docx.
    Lève ProvenanceError si un chiffre non sourcé est détecté.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx est requis pour la génération docx.")

    # Vérification provenance avant génération
    for res in resultats:
        _verifier_sources(res.get("valeur"), res.get("sources", []),
                          f"contrôle {res.get('controle_ref')}")

    # Valeurs calculées par Python (pour vérification des feuilles IA)
    valeurs_python: set[float] = {
        float(r["valeur"]) for r in resultats
        if r.get("valeur") is not None and isinstance(r.get("valeur"), (int, float))
    }
    avertissements_provenance: list[str] = []
    for ft in feuilles:
        contenu = ft.get("contenu_redige", "") or ""
        if contenu:
            warns = _verifier_provenance_texte_llm(contenu, valeurs_python, ft.get("cycle", "?"))
            avertissements_provenance.extend(warns)

    doc = Document()

    # Style du document
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # En-tête
    titre = doc.add_heading("DOSSIER DE TRAVAIL", 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.add_run(f"Client : {projet.get('client', 'N/A')}\n").bold = True
    info.add_run(f"Exercice : {projet.get('exercice', 'N/A')}\n")
    _seuil = projet.get('seuil_signification')
    _seuil_txt = f"{_seuil:,.0f} FDJ" if isinstance(_seuil, (int, float)) else "Non défini"
    info.add_run(f"Seuil de signification : {_seuil_txt}\n")
    info.add_run(f"Généré le : {_now()}\n")
    info.add_run(f"Référence {norme(230)} : Dossier de travail\n")

    doc.add_paragraph("─" * 60)

    # Numérotation dynamique des sections (certaines sont conditionnelles)
    _num_section = 0

    def _titre_section(libelle: str) -> str:
        nonlocal _num_section
        _num_section += 1
        return f"{_num_section}. {libelle}"

    # Section résultats des contrôles
    doc.add_heading(_titre_section("Résultats des contrôles"), level=1)

    for res in resultats:
        sources_str = ", ".join(str(s)[:20] for s in (res.get("sources") or [])[:3])
        if len(res.get("sources") or []) > 3:
            sources_str += "..."

        p = doc.add_paragraph(style="List Bullet")
        label = f"[{res['controle_ref']}] {res.get('details', '')}"
        run = p.add_run(label)
        if res["statut"] == "exception":
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

        if sources_str:
            p.add_run(f"\n   Sources : {sources_str}").font.size = Pt(9)

    # Section exceptions
    doc.add_heading(_titre_section("Exceptions et leur traitement"), level=1)

    open_exc = [e for e in exceptions if e.get("statut") == "ouverte"]
    closed_exc = [e for e in exceptions if e.get("statut") == "tranchee"]

    if open_exc:
        doc.add_paragraph(f"⚠ {len(open_exc)} exception(s) ouverte(s) non tranchée(s).")

    for exc in exceptions:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"[{exc.get('nep_ref')}] {exc.get('controle_ref')} — "
                        f"Sévérité : {exc.get('severite', 'N/A')}")
        if exc.get("statut") == "ouverte":
            run.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
        else:
            run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)

        p.add_run(f"\n   {exc.get('description', '')}")
        if exc.get("decision_humaine"):
            p.add_run(f"\n   Décision : {exc['decision_humaine']} (par {exc.get('decideur', 'N/A')})")
        type_res = exc.get("type_resolution")
        if type_res:
            labels_res = {
                "corrigee": "Anomalie corrigée par le client",
                "sans_incidence": "Sans incidence — explication obtenue, aucune anomalie avérée",
                "non_corrigee": "Anomalie NON corrigée",
            }
            txt_res = labels_res.get(type_res, type_res)
            mi = exc.get("montant_incidence")
            if type_res == "non_corrigee" and isinstance(mi, (int, float)):
                txt_res += f" — incidence : {mi:,.2f}"
            p.add_run(f"\n   Résolution ({norme(450)}) : {txt_res}")
        if exc.get("interpretation_llm"):
            p.add_run(f"\n   Interprétation : {exc['interpretation_llm'][:200]}...")

    # Section synthèse NEP 450 — cumul des anomalies non corrigées vs seuil
    if synthese_anomalies:
        doc.add_heading(_titre_section(f"Synthèse des anomalies ({norme(450)})"), level=1)
        sa = synthese_anomalies
        p_syn = doc.add_paragraph()
        p_syn.add_run(
            f"Anomalies corrigées : {sa.get('nb_corrigees', 0)} — "
            f"Sans incidence : {sa.get('nb_sans_incidence', 0)} — "
            f"Non corrigées : {sa.get('nb_non_corrigees', 0)} — "
            f"Tranchées sans typologie : {sa.get('nb_non_typees', 0)}\n"
        )
        cumul = sa.get("cumul_non_corrigees", 0.0)
        seuil_sig = sa.get("seuil_signification")
        run_cumul = p_syn.add_run(
            f"Cumul des anomalies non corrigées : {cumul:,.2f}"
            + (f" / seuil de signification : {seuil_sig:,.2f}" if seuil_sig else " (seuil non défini)")
        )
        run_cumul.bold = True
        if sa.get("depasse_seuil_signification"):
            p_alerte = doc.add_paragraph()
            r = p_alerte.add_run(
                "⚠ Le cumul des anomalies non corrigées DÉPASSE le seuil de signification. "
                f"Conformément à la norme {norme(450)}, ce dépassement doit être pris en compte dans "
                "la formulation de l'opinion (réserve ou refus de certifier à envisager)."
            )
            r.bold = True
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        elif seuil_sig:
            doc.add_paragraph(
                "Le cumul des anomalies non corrigées reste inférieur au seuil de "
                "signification. Prises isolément et en cumul, elles n'ont pas d'incidence "
                "significative sur les comptes pris dans leur ensemble."
            )
        for e_nc in sa.get("exceptions_non_corrigees", []):
            mi = e_nc.get("montant_incidence")
            doc.add_paragraph(
                f"[{e_nc.get('controle_ref')}] {(e_nc.get('description') or '')[:120]} — "
                f"incidence : {mi:,.2f}" if isinstance(mi, (int, float)) else
                f"[{e_nc.get('controle_ref')}] {(e_nc.get('description') or '')[:120]}",
                style="List Bullet",
            )

    # Section contrôles non exécutés (NEP 230)
    if controles_ignores:
        doc.add_heading(_titre_section(f"Contrôles prévus non exécutés ({norme(230)})"), level=1)
        doc.add_paragraph(
            "Les contrôles suivants n'ont pas pu être exécutés lors de la dernière "
            f"passe ; le motif est documenté conformément à la norme {norme(230)}."
        )
        for ci in controles_ignores:
            doc.add_paragraph(
                f"[{ci.get('controle_ref')}] ({ci.get('cycle', '?')}) — {ci.get('raison', '')}",
                style="List Bullet",
            )

    # Section feuilles de travail
    doc.add_heading(_titre_section("Feuilles de travail par cycle"), level=1)

    for ft in feuilles:
        doc.add_heading(f"Cycle : {ft.get('cycle', 'N/A')}", level=2)
        doc.add_paragraph(ft.get("contenu_redige", ""))
        sources_ft = ft.get("sources", [])
        if sources_ft:
            doc.add_paragraph(f"Sources : {', '.join(str(s) for s in sources_ft[:5])}")

    # Avertissements de provenance (si des nombres LLM ne sont pas tracés)
    if avertissements_provenance:
        doc.add_heading("⚠ Avertissements de traçabilité", level=1)
        p_warn = doc.add_paragraph()
        p_warn.add_run(
            "Les montants suivants ont été détectés dans les feuilles de travail rédigées par l'IA "
            "sans correspondance exacte dans les valeurs calculées par le moteur déterministe. "
            "L'auditeur doit les vérifier manuellement."
        ).italic = True
        for w in avertissements_provenance[:20]:
            doc.add_paragraph(w, style="List Bullet")

    # Pied de page
    doc.add_paragraph("─" * 60)
    footer = doc.add_paragraph()
    footer.add_run("Ce dossier est généré automatiquement par Probare. "
                   "L'opinion d'audit reste de la responsabilité exclusive "
                   "de l'auditeur habilité signataire.")
    footer.runs[0].italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def generer_note_planification(
    projet: dict,
    plan: dict,
    risques: list[dict],
    programme: list[dict],
    note_synthese: dict,
    output_path: Path,
) -> Path:
    """Génère la Note de Planification de l'Audit en .docx (NEP 300).

    Document complet structuré en 7 sections + page de garde,
    destiné au dossier de travail de la mission.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx est requis.")

    VIOLET = RGBColor(0x4F, 0x46, 0xE5)
    ROUGE  = RGBColor(0xC0, 0x39, 0x2B)
    ORANGE = RGBColor(0xE6, 0x7E, 0x22)
    VERT   = RGBColor(0x27, 0xAE, 0x60)

    NIVEAU_LABELS = {
        "eleve":  "ÉLEVÉ",
        "moyen":  "MOYEN",
        "faible": "FAIBLE",
    }

    doc = Document()

    for section in doc.sections:
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def h1(text: str):
        p = doc.add_heading(text, level=1)
        if p.runs:
            p.runs[0].font.color.rgb = VIOLET
        return p

    def h2(text: str):
        return doc.add_heading(text, level=2)

    def para(text: str, bold: bool = False, italic: bool = False, size: int = 11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.size = Pt(size)
        return p

    def sep():
        p = doc.add_paragraph("─" * 80)
        p.paragraph_format.space_after = Pt(0)
        if p.runs:
            p.runs[0].font.size = Pt(8)

    def fmt_montant(v: Any) -> str:
        try:
            return f"{float(v):,.0f} FDJ"
        except Exception:
            return str(v) if v is not None else "—"

    def label_val(p: Any, label: str, val: Any) -> None:
        r1 = p.add_run(f"{label} : ")
        r1.bold = True
        r1.font.size = Pt(11)
        p.add_run(str(val) if val is not None else "—").font.size = Pt(11)
        p.add_run("\n")

    # ── Index des sections IA ────────────────────────────────────────────────
    sections_ia: dict[str, str] = {}
    for s in (note_synthese.get("sections") or []):
        t = (s.get("titre") or "").strip()
        c = (s.get("contenu") or "").strip()
        if t:
            sections_ia[t] = c
    # Raccourcis tolérants
    def ia(key: str) -> str:
        for k, v in sections_ia.items():
            if key.lower() in k.lower():
                return v
        return ""

    client_nom = projet.get("client") or projet.get("nom") or "N/A"
    exercice   = projet.get("exercice") or "N/A"
    seuil_calc = plan.get("seuil_calcule")
    seuil_plan = plan.get("seuil_planification_calcule")
    agregat_type = plan.get("agregat_type") or "—"
    agregat_val  = plan.get("agregat_valeur")
    taux_sig     = plan.get("taux_signification")
    taux_plan    = plan.get("taux_planification")

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # PAGE DE GARDE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    pg = doc.add_paragraph()
    pg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = pg.add_run("PROBARE")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = VIOLET

    pg2 = doc.add_paragraph()
    pg2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pg2.add_run("Logiciel d'audit IA-first — Djibouti").italic = True

    doc.add_paragraph()
    titre = doc.add_heading("NOTE DE PLANIFICATION DE L'AUDIT", level=0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    sep()
    doc.add_paragraph()

    garde = doc.add_paragraph()
    garde.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for lbl, val in [
        ("Entité auditée",    client_nom),
        ("Exercice",          exercice),
        ("Date d'émission",   _now()),
        ("Référentiel",       f"{norme(300)} — Planification"),
        ("Statut",            "Document de travail — CONFIDENTIEL"),
    ]:
        r2 = garde.add_run(f"{lbl} : {val}\n")
        r2.font.size = Pt(12)

    doc.add_paragraph()
    sep()
    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # SOMMAIRE
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1("Sommaire")
    for item in [
        "1. Cadre et objectifs de la mission",
        f"2. Connaissance de l'entité ({norme(315)})",
        f"3. Procédures analytiques préliminaires ({norme(520)})",
        f"4. Seuil de signification ({norme(320)})",
        f"5. Cartographie des risques ({norme(315)})",
        f"6. Programme de travail ({norme(300)} / {norme(330)})",
        "7. Synthèse et conclusion",
    ]:
        doc.add_paragraph(item, style="List Number")
    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 1. CADRE ET OBJECTIFS
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1("1. Cadre et objectifs de la mission")
    para(
        f"La présente note de planification est établie conformément à la norme {norme(300)} "
        f"pour la mission d'audit des comptes de l'exercice {exercice} "
        f"de l'entité {client_nom}. Elle constitue un document de travail confidentiel "
        "intégré au dossier d'audit et doit être conservée pendant la durée légale (5 ans)."
    )
    h2("1.1 Objectifs de l'audit")
    para(
        "L'audit a pour objectif d'émettre une opinion sur la régularité, la sincérité "
        "et l'image fidèle des états financiers. Cet objectif est atteint par la mise en "
        "œuvre de diligences adaptées aux risques identifiés lors de la présente phase "
        "de planification."
    )
    h2("1.2 Référentiels applicables")
    for nep in [
        f"{norme(300)} — Planification de la mission",
        f"{norme(315)} — Connaissance de l'entité et identification des risques",
        f"{norme(320)} — Seuil de signification",
        f"{norme(330)} — Procédures d'audit mises en œuvre en réponse aux risques évalués",
        f"{norme(520)} — Procédures analytiques",
        f"{norme(230)} — Documentation des travaux",
    ]:
        doc.add_paragraph(nep, style="List Bullet")
    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 2. CONNAISSANCE DE L'ENTITÉ
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(f"2. Connaissance de l'entité ({norme(315)})")

    contenu_entite = ia("connaissance")
    if contenu_entite:
        para(contenu_entite)
    doc.add_paragraph()

    h2("2.1 Identification et présentation générale")
    p_gen = doc.add_paragraph()
    label_val(p_gen, "Forme juridique",        plan.get("forme_juridique"))
    label_val(p_gen, "Date de création",       plan.get("date_creation_entreprise"))
    label_val(p_gen, "Effectif",               plan.get("effectif"))
    label_val(p_gen, "Système d'information",  plan.get("systeme_information"))

    if plan.get("activites_principales"):
        h2("2.2 Activités et marchés")
        para(plan["activites_principales"])
        if plan.get("marches_principaux"):
            para(f"Marchés principaux : {plan['marches_principaux']}")

    h2("2.3 Gouvernance — Dirigeants")
    dirigeants = plan.get("dirigeants") or []
    if dirigeants:
        t = doc.add_table(rows=1, cols=3)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, col in enumerate(["Nom", "Fonction", "Email"]):
            hdr[i].text = col
            hdr[i].paragraphs[0].runs[0].bold = True
        for d in dirigeants:
            row = t.add_row().cells
            row[0].text = d.get("nom") or "—"
            row[1].text = d.get("fonction") or "—"
            row[2].text = d.get("email") or "—"
    else:
        para("Aucun dirigeant renseigné.", italic=True)

    h2("2.4 Facteurs de risque inhérent identifiés")
    facteurs = plan.get("facteurs_risque_inherent") or []
    if facteurs:
        for f in facteurs:
            doc.add_paragraph(str(f), style="List Bullet")
    else:
        para("Aucun facteur de risque inhérent spécifique identifié.", italic=True)

    if plan.get("observations"):
        h2("2.5 Observations complémentaires")
        para(plan["observations"])

    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 3. PROCÉDURES ANALYTIQUES
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(f"3. Procédures analytiques préliminaires ({norme(520)})")

    contenu_analytique = ia("analytique") or ia("variations")
    if contenu_analytique:
        para(contenu_analytique)
    else:
        para(
            "Les procédures analytiques préliminaires consistent à comparer les données "
            "financières de l'exercice N aux données de l'exercice N-1 afin d'identifier "
            "les variations significatives susceptibles de révéler un risque d'anomalie."
        )

    h2("3.1 Variations significatives identifiées")
    variations: list[dict] = []
    vraw = plan.get("variations_json")
    if isinstance(vraw, str):
        try:
            variations = json.loads(vraw)
        except Exception:
            pass
    elif isinstance(vraw, list):
        variations = vraw

    sig = [v for v in variations if v.get("significative")]
    if sig:
        t = doc.add_table(rows=1, cols=5)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, col in enumerate(["Compte", "Libellé", "Solde N", "Solde N-1", "Variation %"]):
            hdr[i].text = col
            hdr[i].paragraphs[0].runs[0].bold = True
        for v in sig[:30]:
            row = t.add_row().cells
            row[0].text = str(v.get("compte") or "")
            row[1].text = str(v.get("libelle") or "")
            row[2].text = fmt_montant(v.get("solde_n"))
            row[3].text = fmt_montant(v.get("solde_n1"))
            pct = v.get("delta_pct")
            row[4].text = f"{pct:+.1f} %" if pct is not None else "—"
        doc.add_paragraph()
    else:
        para("Aucune variation significative identifiée dans les données disponibles.", italic=True)

    h2("3.2 Interprétation")
    interp: dict | None = None
    iraw = plan.get("interpretation_variations")
    if isinstance(iraw, str):
        try:
            interp = json.loads(iraw)
        except Exception:
            pass
    elif isinstance(iraw, dict):
        interp = iraw

    if interp:
        if interp.get("synthese"):
            para(interp["synthese"])
        zones = interp.get("zones_risque") or []
        if zones:
            t2 = doc.add_table(rows=1, cols=3)
            t2.style = "Table Grid"
            hdr2 = t2.rows[0].cells
            for i, col in enumerate(["Cycle / Compte", "Niveau", "Explication"]):
                hdr2[i].text = col
                hdr2[i].paragraphs[0].runs[0].bold = True
            for z in zones:
                row = t2.add_row().cells
                row[0].text = z.get("libelle") or z.get("cycle") or "—"
                row[1].text = z.get("niveau") or "—"
                row[2].text = z.get("explication") or "—"
            doc.add_paragraph()
        if interp.get("facteurs_contextuels"):
            para(f"Facteurs contextuels : {interp['facteurs_contextuels']}")
        alertes = interp.get("alertes") or []
        if alertes:
            h2("Points d'attention")
            for a in alertes:
                doc.add_paragraph(str(a), style="List Bullet")
    else:
        para("Interprétation analytique non disponible.", italic=True)

    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 4. SEUIL DE SIGNIFICATION
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(f"4. Seuil de signification ({norme(320)})")

    contenu_seuil = ia("seuil") or ia("risque")
    if contenu_seuil:
        para(contenu_seuil)

    h2("4.1 Base de calcul et paramètres retenus")
    p_seuil = doc.add_paragraph()
    label_val(p_seuil, "Agrégat retenu",          agregat_type)
    label_val(p_seuil, "Valeur de l'agrégat",     fmt_montant(agregat_val))
    label_val(p_seuil, "Taux de signification",   f"{taux_sig} %" if taux_sig else "—")
    label_val(p_seuil, "Taux de planification",   f"{taux_plan} %" if taux_plan else "—")
    label_val(p_seuil, "Seuil de signification",  fmt_montant(seuil_calc))
    label_val(p_seuil, "Seuil de planification",  fmt_montant(seuil_plan))

    h2("4.2 Justification du seuil retenu")
    para(
        f"Le seuil de signification de {fmt_montant(seuil_calc)} ({taux_sig or '—'} % "
        f"de {agregat_type}) est le montant au-delà duquel une anomalie, prise "
        f"isolément ou EN CUMUL avec les autres anomalies non corrigées ({norme(450)}), "
        "est susceptible d'influencer le jugement d'un utilisateur des comptes et "
        "d'affecter l'opinion d'audit. Une anomalie inférieure au seuil ne peut être "
        "considérée comme sans incidence qu'après agrégation avec l'ensemble des "
        "anomalies non corrigées relevées au cours de la mission. "
        "Le seuil de planification de "
        f"{fmt_montant(seuil_plan)} ({taux_plan or '—'} %) est appliqué pour les "
        "tests de détail afin de conserver une marge de sécurité permettant d'absorber "
        "les anomalies non détectées. Ces paramètres sont conformes aux pratiques "
        f"professionnelles et aux exigences de la norme {norme(320)}."
    )
    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 5. CARTOGRAPHIE DES RISQUES
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(f"5. Cartographie des risques ({norme(315)})")
    para(f"{len(risques)} risque(s) identifié(s) et validé(s) pour cette mission.")

    eleve  = [r for r in risques if r.get("niveau") == "eleve"]
    moyen  = [r for r in risques if r.get("niveau") == "moyen"]
    faible = [r for r in risques if r.get("niveau") == "faible"]
    para(f"Dont : {len(eleve)} élevé(s), {len(moyen)} moyen(s), {len(faible)} faible(s).")
    doc.add_paragraph()

    cycles_risques: dict[str, list[dict]] = {}
    for r in risques:
        c = r.get("cycle") or "transversal"
        cycles_risques.setdefault(c, []).append(r)

    for cycle_id, cycle_risques in sorted(cycles_risques.items()):
        h2(f"Cycle : {cycle_id.capitalize()}")
        t = doc.add_table(rows=1, cols=5)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, col in enumerate(["Libellé du risque", "Niveau", "Source", "Assertions ISA", "Commentaire"]):
            hdr[i].text = col
            hdr[i].paragraphs[0].runs[0].bold = True
        for r in cycle_risques:
            row = t.add_row().cells
            row[0].text = r.get("libelle") or "—"
            row[1].text = NIVEAU_LABELS.get(r.get("niveau", ""), r.get("niveau", "—"))
            row[2].text = "IA" if r.get("issu_ia") else "Manuel"
            assertions = r.get("assertions") or []
            row[3].text = ", ".join(assertions) if assertions else "—"
            row[4].text = (r.get("commentaire") or "")[:100]
        doc.add_paragraph()

    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 6. PROGRAMME DE TRAVAIL
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1(f"6. Programme de travail ({norme(300)} / {norme(330)})")

    contenu_programme = ia("programme") or ia("justification")
    if contenu_programme:
        para(contenu_programme)

    inclus = [p for p in programme if p.get("statut") == "inclus"]
    exclus = [p for p in programme if p.get("statut") != "inclus"]
    para(
        f"{len(inclus)} contrôle(s) inclus dans le programme sur {len(programme)} planifiés. "
        f"{len(exclus)} contrôle(s) exclu(s) sur la base du niveau de risque et du seuil de signification."
    )

    cycles_prog: dict[str, list[dict]] = {}
    for item in inclus:
        c = item.get("cycle") or "transversal"
        cycles_prog.setdefault(c, []).append(item)

    for cycle_id, items in sorted(cycles_prog.items()):
        h2(f"Cycle : {cycle_id.capitalize()}")
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        hdr = t.rows[0].cells
        for i, col in enumerate(["Réf. contrôle", "Libellé", "Priorité", "Notes"]):
            hdr[i].text = col
            hdr[i].paragraphs[0].runs[0].bold = True
        for item in items:
            row = t.add_row().cells
            row[0].text = item.get("controle_ref") or "—"
            row[1].text = (item.get("libelle") or "")[:120]
            row[2].text = item.get("priorite") or "—"
            row[3].text = (item.get("notes") or "")[:80]
        doc.add_paragraph()

    doc.add_page_break()

    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    # 7. SYNTHÈSE ET CONCLUSION
    # ━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
    h1("7. Synthèse et conclusion")

    if note_synthese.get("titre"):
        para(note_synthese["titre"], bold=True)
        doc.add_paragraph()

    for s in (note_synthese.get("sections") or []):
        h2(s.get("titre") or "")
        para(s.get("contenu") or "")

    if note_synthese.get("conclusion"):
        doc.add_paragraph()
        p_conc = doc.add_paragraph()
        r_conc = p_conc.add_run(note_synthese["conclusion"])
        r_conc.italic = True

    doc.add_paragraph()
    sep()

    # Zone de signature
    doc.add_paragraph()
    sig_p = doc.add_paragraph()
    sig_p.add_run("Validation par l'auditeur responsable de la mission\n\n").bold = True
    sig_p.add_run("Nom : ___________________________          ")
    sig_p.add_run("Date : ___________________________\n\n")
    sig_p.add_run("Signature : ___________________________")

    doc.add_paragraph()
    sep()

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = footer_p.add_run(
        f"Document généré par Probare le {_now()} · "
        f"Document confidentiel · Dossier de travail {norme(230)}"
    )
    r_f.font.size = Pt(9)
    r_f.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def generer_tableau_exceptions(
    exceptions: list[dict],
    output_path: Path,
) -> Path:
    """Génère un tableau Excel des exceptions."""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        raise ImportError("openpyxl est requis pour la génération xlsx.")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Exceptions"

    headers = ["ID", "Contrôle", "NEP", "Sévérité", "Description",
               "Statut", "Décision", "Décideur", "Horodatage"]
    header_fill = PatternFill("solid", fgColor="4F46E5")
    header_font = Font(bold=True, color="FFFFFF")

    for col_idx, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col_idx, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")

    for row_idx, exc in enumerate(exceptions, 2):
        ws.cell(row=row_idx, column=1, value=exc.get("id", "")[:8])
        ws.cell(row=row_idx, column=2, value=exc.get("controle_ref", ""))
        ws.cell(row=row_idx, column=3, value=exc.get("nep_ref", ""))
        ws.cell(row=row_idx, column=4, value=exc.get("severite", ""))
        ws.cell(row=row_idx, column=5, value=exc.get("description", "")[:200])
        ws.cell(row=row_idx, column=6, value=exc.get("statut", ""))
        ws.cell(row=row_idx, column=7, value=exc.get("decision_humaine", ""))
        ws.cell(row=row_idx, column=8, value=exc.get("decideur", ""))
        ws.cell(row=row_idx, column=9, value=exc.get("horodatage", ""))

        if exc.get("statut") == "ouverte":
            fill = PatternFill("solid", fgColor="FFF3CD")
        else:
            fill = PatternFill("solid", fgColor="D4EDDA")
        for col in range(1, len(headers) + 1):
            ws.cell(row=row_idx, column=col).fill = fill

    for col in ws.columns:
        ws.column_dimensions[col[0].column_letter].auto_size = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(output_path))
    return output_path


# ─── Demande de diligences au client (#9) ─────────────────────────────────────

# Cycle (préfixe de la référence de contrôle) → libellé lisible
_CYCLE_LABELS = {
    "TRESOR": "Trésorerie", "ACHAT": "Achats-Fournisseurs", "VENTE": "Ventes-Clients",
    "IMO": "Immobilisations", "STOCK": "Stocks", "PAIE": "Personnel-Paie",
    "TAXE": "Impôts et taxes", "CP": "Capitaux propres et provisions",
}


def _cycle_depuis_ref(ref: str) -> str:
    prefixe = (ref or "").split("-")[0]
    return _CYCLE_LABELS.get(prefixe, "Divers")


_SEVERITE_LABELS = {
    "critique": "Critique", "significative": "Significative", "mineure": "Mineure",
}


_SOURCE_LABELS = {
    "grand_livre": "Grand livre",
    "balance": "Balance générale",
    "releve_bancaire": "Relevé bancaire",
    "annexe": "Document annexe",
}


def _libelle_source(fichier: dict) -> str:
    """Libellé lisible d'une source de détection (type de document + nom)."""
    type_doc = fichier.get("type_document") or fichier.get("type") or ""
    label = _SOURCE_LABELS.get(type_doc)
    nom = fichier.get("nom") or ""
    if label:
        return f"{label} ({nom})" if nom else label
    return nom or "Source comptable"


def generer_demande_diligences(
    projet: dict,
    exceptions: list[dict],
    output_path: Path,
    seulement_ouvertes: bool = True,
    fichiers_map: dict | None = None,
) -> Path:
    """Génère une demande de diligences .docx présentable au client (#9).

    Reprend chaque exception (anomalie relevée), regroupée par cycle, avec sa
    description, la source de détection (grand livre, balance, relevé…), les
    hypothèses de cause et les diligences/pièces à fournir.
    Mise en page professionnelle, prête à imprimer/envoyer.

    `fichiers_map` : {fichier_source_id: fichier_source_dict} pour résoudre la
    source de détection de chaque exception (#3).
    """
    fichiers_map = fichiers_map or {}
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx est requis pour la génération docx.")

    VIOLET = RGBColor(0x4F, 0x46, 0xE5)
    exceptions = [e for e in exceptions
                  if not seulement_ouvertes or e.get("statut") == "ouverte"]

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    titre = doc.add_heading("DEMANDE DE DILIGENCES", 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Points relevés lors de l'audit — pièces et explications attendues").italic = True

    info = doc.add_paragraph()
    info.add_run(f"Client : {projet.get('client', 'N/A')}\n").bold = True
    info.add_run(f"Exercice : {projet.get('exercice', 'N/A')}\n")
    info.add_run(f"Édité le : {_now()}\n")
    info.add_run(f"Nombre de points : {len(exceptions)}\n")

    doc.add_paragraph("─" * 60)

    intro = doc.add_paragraph()
    intro.add_run(
        "Dans le cadre de nos travaux d'audit, les points suivants nécessitent votre "
        "attention. Pour chacun, merci de nous communiquer les explications et les pièces "
        "justificatives correspondantes. Ces éléments nous permettront de conclure sur les "
        "zones concernées. Sauf mention contraire, tous les montants sont exprimés en FDJ."
    ).italic = True

    if not exceptions:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Aucun point en attente à ce jour.").bold = True
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return output_path

    # Regroupement par cycle
    par_cycle: dict[str, list[dict]] = {}
    for e in exceptions:
        par_cycle.setdefault(_cycle_depuis_ref(e.get("controle_ref", "")), []).append(e)

    compteur = 0
    for cycle in sorted(par_cycle):
        doc.add_heading(f"Cycle : {cycle}", level=1)
        for exc in par_cycle[cycle]:
            compteur += 1
            h = doc.add_heading(level=2)
            sev = _SEVERITE_LABELS.get(exc.get("severite", ""), exc.get("severite", ""))
            run = h.add_run(f"Point {compteur}"
                            + (f" — {sev}" if sev else ""))
            run.font.color.rgb = VIOLET

            p = doc.add_paragraph()
            p.add_run("Constat : ").bold = True
            p.add_run(exc.get("description", "") or "—")

            # Source de détection (#3) : grand livre, balance, relevé bancaire…
            sources_ids = exc.get("fichiers_sources") or []
            labels_sources = []
            for fid in sources_ids:
                fichier = fichiers_map.get(fid)
                if fichier:
                    lib = _libelle_source(fichier)
                    if lib not in labels_sources:
                        labels_sources.append(lib)
            if labels_sources:
                p_src = doc.add_paragraph()
                p_src.add_run("Source de détection : ").bold = True
                p_src.add_run(", ".join(labels_sources))

            explication = exc.get("interpretation_llm") or exc.get("explication")
            if explication:
                p = doc.add_paragraph()
                p.add_run("Analyse préliminaire : ").bold = True
                p.add_run(str(explication))

            hypotheses = exc.get("hypotheses") or []
            if hypotheses:
                doc.add_paragraph("Causes possibles à confirmer :").runs[0].bold = True
                for hyp in hypotheses:
                    doc.add_paragraph(str(hyp), style="List Bullet")

            diligences = exc.get("diligences") or []
            if diligences:
                doc.add_paragraph("Éléments et pièces attendus de votre part :").runs[0].bold = True
                for dil in diligences:
                    doc.add_paragraph(str(dil), style="List Bullet")

            # Espace réponse client
            rep = doc.add_paragraph()
            rep.add_run("Réponse / pièces fournies : ").italic = True
            rep.add_run("_" * 55)
            doc.add_paragraph()

    doc.add_paragraph("─" * 60)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"Document généré par Probare le {_now()} · Confidentiel")
    r.font.size = Pt(9)
    r.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ─── Questionnaire de contrôle interne vierge, à imprimer (#2) ─────────────────

_QCI_CYCLE_LABELS = {
    "tresorerie": "Trésorerie", "achats": "Achats-Fournisseurs", "ventes": "Ventes-Clients",
    "immobilisations": "Immobilisations", "stocks": "Stocks", "paie": "Personnel-Paie",
    "impots": "Impôts et taxes", "capitaux_propres": "Capitaux propres et provisions",
}


def generer_questionnaire_vierge(
    projet: dict,
    qci_par_cycle: dict,
    cycles: list[str],
    output_path: Path,
) -> Path:
    """Génère le questionnaire de contrôle interne VIERGE en .docx, prêt à imprimer
    et à faire remplir sur le terrain (#2). Colonnes Oui / Non / N.A. + commentaire."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx est requis pour la génération docx.")

    VIOLET = RGBColor(0x4F, 0x46, 0xE5)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    titre = doc.add_heading("QUESTIONNAIRE DE CONTRÔLE INTERNE", 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.add_run(f"Client : {projet.get('client', 'N/A')}").bold = True
    info.add_run(f"          Exercice : {projet.get('exercice', 'N/A')}\n")
    info.add_run("Rempli par : ____________________     Fonction : ____________________     "
                 "Date : ____________")

    doc.add_paragraph(
        "Pour chaque question, cochez Oui, Non ou N.A. (non applicable) et précisez un "
        "commentaire si nécessaire.", style=None
    ).italic = True

    cycles = [c for c in cycles if c in qci_par_cycle] or list(qci_par_cycle.keys())
    for cycle in cycles:
        questions = qci_par_cycle.get(cycle, [])
        if not questions:
            continue
        h = doc.add_heading(level=1)
        h.add_run(f"Cycle : {_QCI_CYCLE_LABELS.get(cycle, cycle.capitalize())}").font.color.rgb = VIOLET

        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        widths = ["N°", "Question", "Oui", "Non", "N.A."]
        hdr = table.rows[0].cells
        for i, col in enumerate(widths):
            hdr[i].text = col
            if hdr[i].paragraphs[0].runs:
                hdr[i].paragraphs[0].runs[0].bold = True
        for idx, q in enumerate(questions, 1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = q.get("question", "")
            row[2].text = "☐"
            row[3].text = "☐"
            row[4].text = "☐"
        # Ligne de commentaires par cycle
        c = doc.add_paragraph()
        c.add_run("Commentaires / observations : ").italic = True
        c.add_run("_" * 70)
        doc.add_paragraph()

    doc.add_paragraph("─" * 60)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"Questionnaire généré par Probare le {_now()} · À remplir sur le terrain")
    r.font.size = Pt(9)
    r.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ─── Synthèse de l'évaluation du contrôle interne, en .docx mis en forme (#2) ──

_NIVEAU_CI_LABELS = {"eleve": "Élevé", "moyen": "Moyen", "faible": "Faible"}
_NIVEAU_CI_COULEURS = {
    "eleve": (0xC0, 0x39, 0x2B),   # rouge
    "moyen": (0xE6, 0x7E, 0x22),   # orange
    "faible": (0x27, 0xAE, 0x60),  # vert
}


def generer_synthese_ci_docx(
    projet: dict,
    synthese: dict,
    output_path: Path,
) -> Path:
    """Génère la synthèse de l'évaluation du contrôle interne en .docx mis en forme (#2).

    `synthese` reprend la structure persistée : titre, sections rédigées,
    conclusion, niveau/score global et matrice des risques par cycle.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx est requis pour la génération docx.")

    VIOLET = RGBColor(0x4F, 0x46, 0xE5)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    titre_txt = synthese.get("titre") or "Synthèse de l'évaluation du contrôle interne"
    titre = doc.add_heading(titre_txt, 0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_paragraph()
    info.add_run(f"Client : {projet.get('client', 'N/A')}\n").bold = True
    info.add_run(f"Exercice : {projet.get('exercice', 'N/A')}\n")
    info.add_run(f"Édité le : {_now()}\n")
    info.add_run(f"Référence {norme(315)} / {norme(330)} : Évaluation du contrôle interne\n")

    doc.add_paragraph("─" * 60)

    # Bandeau appréciation globale
    niveau = synthese.get("niveau_global") or ""
    score = synthese.get("score_global")
    p_glob = doc.add_paragraph()
    p_glob.add_run("Appréciation globale du contrôle interne : ").bold = True
    run_niv = p_glob.add_run(_NIVEAU_CI_LABELS.get(niveau, niveau or "N/A"))
    run_niv.bold = True
    coul = _NIVEAU_CI_COULEURS.get(niveau)
    if coul:
        run_niv.font.color.rgb = RGBColor(*coul)
    if isinstance(score, (int, float)):
        p_glob.add_run(f"   (score moyen : {score:.2f})")
    nb = synthese.get("nb_cycles_evalues")
    if nb:
        p_glob.add_run(f"   ·   {nb} cycle(s) évalué(s)")

    # Matrice des risques par cycle
    matrice = synthese.get("matrice") or []
    if matrice:
        doc.add_heading("Matrice des risques par cycle", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, col in enumerate(["Cycle", "Niveau de risque", "Score"]):
            hdr[i].text = col
            if hdr[i].paragraphs[0].runs:
                hdr[i].paragraphs[0].runs[0].bold = True
        for m in matrice:
            cyc = m.get("cycle") or ""
            row = table.add_row().cells
            row[0].text = _QCI_CYCLE_LABELS.get(cyc, cyc.capitalize() if cyc else "—")
            niv_c = m.get("niveau_risque") or ""
            row[1].text = _NIVEAU_CI_LABELS.get(niv_c, niv_c or "—")
            if row[1].paragraphs[0].runs:
                coul_c = _NIVEAU_CI_COULEURS.get(niv_c)
                if coul_c:
                    row[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(*coul_c)
            sc = m.get("score")
            row[2].text = f"{sc:.2f}" if isinstance(sc, (int, float)) else "—"
        doc.add_paragraph()

    # Sections rédigées
    for sec in synthese.get("sections") or []:
        titre_sec = sec.get("titre") or ""
        contenu = sec.get("contenu") or ""
        if titre_sec:
            h = doc.add_heading(level=1)
            h.add_run(titre_sec).font.color.rgb = VIOLET
        if contenu:
            doc.add_paragraph(str(contenu))

    # Conclusion
    conclusion = synthese.get("conclusion")
    if conclusion:
        doc.add_heading("Conclusion", level=1)
        p_ccl = doc.add_paragraph()
        p_ccl.add_run(str(conclusion))

    doc.add_paragraph("─" * 60)
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"Synthèse générée par Probare le {_now()} · Confidentiel")
    r.font.size = Pt(9)
    r.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


# ─── Helpers cabinet (identité signataire, partagés par les livrables finaux) ──


def _cabinet_ligne_adresse(cabinet: dict) -> str:
    parts = [
        cabinet.get("adresse_rue"),
        " ".join(p for p in [cabinet.get("adresse_code_postal"),
                             cabinet.get("adresse_ville")] if p),
        cabinet.get("adresse_pays"),
    ]
    return " — ".join(p for p in parts if p)


def _entete_cabinet(doc, cabinet: dict) -> None:
    """Rend l'en-tête d'identité du cabinet en haut d'un livrable signé.
    `cabinet` provient du paramétrage Cabinet (transmis par le frontend)."""
    from docx.shared import Pt, RGBColor
    cabinet = cabinet or {}
    if not any(cabinet.get(k) for k in ("nom", "responsable_nom", "adresse_rue", "email")):
        return
    p = doc.add_paragraph()
    if cabinet.get("nom"):
        r = p.add_run(cabinet["nom"])
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    if cabinet.get("forme_juridique"):
        p.add_run(f"  ·  {cabinet['forme_juridique']}").font.size = Pt(9)
    ligne_adr = _cabinet_ligne_adresse(cabinet)
    contacts = "  ·  ".join(x for x in [
        cabinet.get("telephone") and f"Tél. {cabinet['telephone']}",
        cabinet.get("email"),
        cabinet.get("site_web"),
    ] if x)
    refs = "  ·  ".join(x for x in [
        cabinet.get("numero_agrement") and f"Agrément {cabinet['numero_agrement']}",
        cabinet.get("numero_ordre") and f"Ordre {cabinet['numero_ordre']}",
    ] if x)
    for txt in (ligne_adr, contacts, refs):
        if txt:
            sub = doc.add_paragraph()
            run = sub.add_run(txt)
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_paragraph("─" * 60)


def _bloc_signature(doc, cabinet: dict) -> None:
    """Rend le bloc « Lieu, date et signature » d'un livrable signé."""
    from docx.shared import Pt
    cabinet = cabinet or {}
    ville = cabinet.get("adresse_ville") or "Djibouti"
    doc.add_paragraph()
    p_lieu = doc.add_paragraph()
    p_lieu.add_run(f"Fait à {ville}, le {datetime.now(timezone.utc).strftime('%d/%m/%Y')}.")
    doc.add_paragraph()
    sig = doc.add_paragraph()
    if cabinet.get("nom"):
        sig.add_run(cabinet["nom"] + "\n").bold = True
    if cabinet.get("responsable_nom"):
        sig.add_run(cabinet["responsable_nom"] + "\n").bold = True
    if cabinet.get("responsable_titre"):
        r = sig.add_run(cabinet["responsable_titre"])
        r.italic = True
    if not cabinet.get("responsable_nom"):
        sig.add_run("\n\nNom et signature : ______________________________")


_MEMO_CYCLE_LABELS = {
    "tresorerie": "Trésorerie", "achats": "Achats-Fournisseurs", "ventes": "Ventes-Clients",
    "immobilisations": "Immobilisations", "stocks": "Stocks", "paie": "Personnel-Paie",
    "impots": "Impôts et taxes", "capitaux_propres": "Capitaux propres et provisions",
}


def _titre_mission(projet: dict) -> str:
    nature = (projet.get("nature_mission") or "contractuelle").lower()
    if "legal" in nature or "légal" in nature or "commissariat" in nature:
        return "RAPPORT DU COMMISSAIRE AUX COMPTES SUR LES COMPTES ANNUELS"
    return "RAPPORT D'AUDIT SUR LES COMPTES ANNUELS"


_TYPE_OPINION_LABELS = {
    "sans_reserve": "Opinion sans réserve",
    "avec_reserve": "Opinion avec réserve",
    "defavorable": "Opinion défavorable",
    "impossibilite": "Impossibilité d'exprimer une opinion",
}


def generer_rapport_audit(
    projet: dict,
    opinion: dict,
    output_path: Path,
    cabinet: dict | None = None,
    plan: dict | None = None,
) -> Path:
    """Génère le RAPPORT D'AUDIT sur les comptes annuels en .docx (ISA/NEP 700).

    Le corps de l'opinion, le fondement et les éventuelles observations sont ceux
    validés par l'auditeur (proposés par l'IA puis éventuellement corrigés).
    Probare ne signe pas : le bloc signature reprend l'identité du cabinet mais
    l'engagement reste celui de l'auditeur habilité.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx est requis pour la génération docx.")

    cabinet = cabinet or {}
    plan = plan or {}
    ref_compta = libelle_referentiel_comptable(projet.get("referentiel_comptable"))
    client = projet.get("client") or projet.get("nom") or "l'entité"
    exercice = projet.get("exercice") or "N"

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _entete_cabinet(doc, cabinet)

    titre = doc.add_heading(_titre_mission(projet), level=0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run(f"Exercice clos — {exercice}").italic = True

    # Destinataire
    dirigeants = plan.get("dirigeants") or []
    destinataire = None
    for d in dirigeants:
        if d.get("fonction") and d.get("nom"):
            destinataire = f"À l'attention de {d['nom']}, {d['fonction']}"
            break
    doc.add_paragraph()
    doc.add_paragraph(destinataire or f"À la direction de {client},")

    # 1. Opinion
    type_op = opinion.get("type_opinion") or "sans_reserve"
    titre_op = opinion.get("titre") or _TYPE_OPINION_LABELS.get(type_op, "Opinion")
    h = doc.add_heading(titre_op, level=1)
    if h.runs:
        h.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    doc.add_paragraph(opinion.get("texte_opinion") or "")

    # 2. Fondement de l'opinion
    doc.add_heading("Fondement de l'opinion", level=1)
    fondement = opinion.get("fondement") or (
        f"Nous avons effectué notre audit selon les normes d'audit {prefixe_actif()}. "
        "Les responsabilités qui nous incombent en vertu de ces normes sont décrites "
        "dans la section « Responsabilités de l'auditeur » du présent rapport. Nous sommes "
        "indépendants de l'entité et estimons que les éléments probants que nous avons "
        "collectés sont suffisants et appropriés pour fonder notre opinion."
    )
    doc.add_paragraph(fondement)

    # 3. Observation / incertitude (conditionnel)
    observations = (opinion.get("observations") or "").strip()
    if observations:
        doc.add_heading("Observation", level=1)
        doc.add_paragraph(observations)

    # 4. Responsabilités de la direction
    doc.add_heading("Responsabilités de la direction relatives aux comptes annuels", level=1)
    doc.add_paragraph(
        f"Il appartient à la direction d'établir des comptes annuels présentant une image "
        f"fidèle conformément à {ref_compta}, ainsi que de mettre en place le contrôle interne "
        "qu'elle estime nécessaire à l'établissement de comptes annuels ne comportant pas "
        "d'anomalies significatives, que celles-ci proviennent de fraudes ou résultent d'erreurs. "
        "Lors de l'établissement des comptes annuels, il incombe à la direction d'évaluer la "
        "capacité de l'entité à poursuivre son exploitation et d'appliquer la convention "
        "comptable de continuité d'exploitation, sauf s'il est prévu de liquider l'entité ou de "
        "cesser son activité."
    )

    # 5. Responsabilités de l'auditeur
    doc.add_heading("Responsabilités de l'auditeur relatives à l'audit des comptes annuels", level=1)
    doc.add_paragraph(
        "Notre objectif est d'obtenir l'assurance raisonnable que les comptes annuels pris dans "
        "leur ensemble ne comportent pas d'anomalies significatives. L'assurance raisonnable "
        "correspond à un niveau élevé d'assurance, sans toutefois garantir qu'un audit réalisé "
        f"conformément aux normes {prefixe_actif()} permette systématiquement de détecter toute "
        "anomalie significative. Dans le cadre de l'audit, nous exerçons notre jugement "
        "professionnel et faisons preuve d'esprit critique : nous identifions et évaluons les "
        "risques d'anomalies significatives, définissons et mettons en œuvre des procédures "
        "d'audit en réponse à ces risques, apprécions le caractère approprié des méthodes "
        "comptables retenues et le caractère raisonnable des estimations, et concluons sur "
        "l'application de la convention de continuité d'exploitation."
    )

    # Signature
    doc.add_paragraph("─" * 60)
    _bloc_signature(doc, cabinet)

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(
        f"Rapport préparé avec l'assistance de Probare le {_now()}. "
        f"L'opinion et la signature relèvent de la responsabilité exclusive de l'auditeur habilité. "
        f"Référentiel d'audit : {prefixe_actif()} · Référentiel comptable : {ref_compta}."
    )
    r.font.size = Pt(8)
    r.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def generer_memorandum_controle_comptes(
    projet: dict,
    resultats: list[dict],
    exceptions: list[dict],
    feuilles: list[dict],
    output_path: Path,
    plan: dict | None = None,
    circularisations: list[dict] | None = None,
    sondages: list[dict] | None = None,
    controles_ignores: list[dict] | None = None,
    cabinet: dict | None = None,
) -> Path:
    """Génère le MÉMORANDUM SUR LE CONTRÔLE DES COMPTES en .docx.

    Structure de référence (rapport bailleur / CAC) : contexte & objectifs, puis
    un chapitre par cycle organisé en triptyque Objectifs → Travaux effectués →
    Commentaires de l'auditeur, et enfin les contrôles prévus non exécutés (NEP 230).
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx est requis pour la génération docx.")

    from ..controls.registry import REGISTRE

    cabinet = cabinet or {}
    plan = plan or {}
    circularisations = circularisations or []
    sondages = sondages or []
    controles_ignores = controles_ignores or []

    ROUGE = RGBColor(0xC0, 0x39, 0x2B)

    client = projet.get("client") or projet.get("nom") or "l'entité"
    exercice = projet.get("exercice") or "N"
    cycles = projet.get("cycles_couverts") or ["tresorerie", "achats", "ventes"]

    # Index par cycle
    feuille_par_cycle = {f.get("cycle"): f for f in feuilles}
    refs_par_cycle: dict[str, set] = {}
    for ref, defn in REGISTRE.items():
        refs_par_cycle.setdefault(defn.cycle, set()).add(ref)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _entete_cabinet(doc, cabinet)

    titre = doc.add_heading("MÉMORANDUM SUR LE CONTRÔLE DES COMPTES", level=0)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.add_run(f"{client} — Exercice clos {exercice}").italic = True

    info = doc.add_paragraph()
    info.add_run(f"Édité le : {_now()}\n").font.size = Pt(9)
    info.add_run(f"Référentiel d'audit : {prefixe_actif()}  ·  "
                 f"Référentiel comptable : {libelle_referentiel_comptable(projet.get('referentiel_comptable'))}"
                 ).font.size = Pt(9)
    doc.add_paragraph("─" * 60)

    # A. Contexte et objectifs
    doc.add_heading("A. Rappel du contexte et des objectifs de la mission", level=1)
    if plan.get("activites_principales"):
        doc.add_paragraph(f"Activité de l'entité : {plan['activites_principales']}")
    if plan.get("forme_juridique"):
        doc.add_paragraph(f"Forme juridique : {plan['forme_juridique']}")
    doc.add_paragraph(
        "Le présent mémorandum a pour objet de fournir des informations complémentaires sur "
        "les rubriques des comptes et sur les principaux travaux que nous avons effectués dans "
        "le cadre de notre examen. Il inclut nos principales conclusions. Notre approche a "
        "consisté à vérifier la réalité, l'exhaustivité et la sincérité des soldes et opérations "
        "de chaque cycle, au moyen des contrôles déterministes du programme de travail, complétés "
        "le cas échéant de confirmations externes et de sondages sur pièces."
    )

    # B. Un chapitre par cycle : Objectifs → Travaux effectués → Commentaires
    lettre = ord("B")
    for cycle in cycles:
        libelle_cycle = _MEMO_CYCLE_LABELS.get(cycle, cycle.capitalize())
        refs = refs_par_cycle.get(cycle, set())
        res_cycle = [r for r in resultats if r.get("controle_ref") in refs]
        exc_cycle = [e for e in exceptions if e.get("controle_ref") in refs]
        circ_cycle = [c for c in circularisations if c.get("cycle") == cycle]
        sond_cycle = [s for s in sondages if s.get("cycle") == cycle]
        feuille = feuille_par_cycle.get(cycle)

        # Ne pas générer de chapitre vide (aucun travail sur ce cycle)
        if not (res_cycle or exc_cycle or circ_cycle or sond_cycle or feuille):
            continue

        doc.add_heading(f"{chr(lettre)}. Cycle {libelle_cycle}", level=1)
        lettre += 1

        # 1. Objectifs
        doc.add_heading("Objectifs", level=2)
        doc.add_paragraph(
            f"Nos travaux sur le cycle {libelle_cycle.lower()} visent à nous assurer de la "
            "réalité, de l'exhaustivité, de l'exactitude et du correct rattachement des soldes "
            "et opérations concernés, au regard des assertions d'audit applicables. À cet effet, "
            "les contrôles suivants ont été prévus :"
        )
        controles_cycle = [defn for ref, defn in REGISTRE.items() if defn.cycle == cycle]
        for defn in controles_cycle:
            doc.add_paragraph(f"{defn.libelle} — {defn.description}", style="List Bullet")

        # 2. Travaux effectués
        doc.add_heading("Travaux effectués", level=2)
        if res_cycle:
            doc.add_paragraph("Contrôles déterministes exécutés :")
            for r in res_cycle:
                statut = "sans anomalie" if r.get("statut") == "ok" else "EXCEPTION"
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"[{r.get('controle_ref')}] ")
                p.add_run(f"{r.get('details', '') or ''} — {statut}")
        if circ_cycle:
            doc.add_paragraph("Confirmations externes (circularisation) :")
            for c in circ_cycle:
                doc.add_paragraph(
                    f"{c.get('libelle') or c.get('compte')} — statut : {c.get('statut', '—')}",
                    style="List Bullet")
        if sond_cycle:
            doc.add_paragraph("Sondages sur pièces :")
            for s in sond_cycle:
                doc.add_paragraph(
                    f"{s.get('libelle') or 'Sondage'} — échantillon de "
                    f"{s.get('taille_echantillon', '—')} sur {s.get('population', '—')} éléments, "
                    f"{s.get('nb_anomalies', 0)} anomalie(s) relevée(s)",
                    style="List Bullet")
        if not (res_cycle or circ_cycle or sond_cycle):
            doc.add_paragraph("Aucun contrôle exécuté sur ce cycle à la date du présent mémorandum.",
                              style="List Bullet")

        # 3. Commentaires de l'auditeur
        doc.add_heading("Commentaires de l'auditeur", level=2)
        if feuille and feuille.get("contenu_redige"):
            doc.add_paragraph(feuille["contenu_redige"])
        if exc_cycle:
            doc.add_paragraph("Anomalies relevées et leur traitement :")
            for e in exc_cycle:
                p = doc.add_paragraph(style="List Bullet")
                run = p.add_run(f"[{e.get('controle_ref')}] {e.get('description', '') or ''}")
                if e.get("statut") == "ouverte":
                    run.font.color.rgb = ROUGE
                type_res = e.get("type_resolution")
                if type_res:
                    labels = {"corrigee": "corrigée par le client",
                              "sans_incidence": "sans incidence (explication obtenue)",
                              "non_corrigee": "non corrigée"}
                    txt = labels.get(type_res, type_res)
                    mi = e.get("montant_incidence")
                    if type_res == "non_corrigee" and isinstance(mi, (int, float)):
                        txt += f" — incidence : {mi:,.0f} FDJ"
                    p.add_run(f" → {txt}")
                elif e.get("statut") == "ouverte":
                    p.add_run(" → en cours d'instruction")
        elif not (feuille and feuille.get("contenu_redige")):
            doc.add_paragraph(
                "Les travaux réalisés sur ce cycle n'appellent pas d'observation particulière "
                "à la date du présent mémorandum."
            )

    # Contrôles non exécutés (NEP 230)
    if controles_ignores:
        doc.add_heading(f"Annexe — Contrôles prévus non exécutés ({norme(230)})", level=1)
        doc.add_paragraph(
            "Les contrôles suivants, prévus au programme de travail, n'ont pas été exécutés ; "
            f"le motif est documenté conformément à la norme {norme(230)}."
        )
        for ci in controles_ignores:
            doc.add_paragraph(
                f"[{ci.get('controle_ref')}] ({ci.get('cycle', '?')}) — {ci.get('raison', '')}",
                style="List Bullet")

    doc.add_paragraph("─" * 60)
    _bloc_signature(doc, cabinet)
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run(f"Mémorandum généré par Probare le {_now()} · Confidentiel · "
                       f"Document de travail {norme(230)}")
    r.font.size = Pt(8)
    r.italic = True

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
