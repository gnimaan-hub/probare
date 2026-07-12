"""Système de mise en forme .docx « Probare » — reproduction fidèle de la
maquette bleu marine / vert du modèle de rapport fourni.

Ce module n'a AUCUNE logique métier : uniquement des primitives de mise en page
(page de garde, bandeaux de section, sous-titres, encadrés de citation, filets
verts, sommaire, pied de page numéroté). Les générateurs de livrables
(`export.py`) l'utilisent pour un rendu homogène.

Couleurs de marque (échantillonnées sur le modèle) :
- NAVY  #052D62  — fonds de titres, bandeaux de section, texte de titres
- GREEN #169D53  — filets, soulignements, accents, marques de citation
"""
from __future__ import annotations

from docx.shared import Pt, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

# ─── Palette ──────────────────────────────────────────────────────────────────
NAVY = RGBColor(0x05, 0x2D, 0x62)
GREEN = RGBColor(0x16, 0x9D, 0x53)
INK = RGBColor(0x1F, 0x2A, 0x37)   # texte courant (gris-anthracite doux)
MUTED = RGBColor(0x64, 0x74, 0x8B)  # légendes / méta
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT = RGBColor(0xF1, 0xF5, 0xF9)  # fonds très clairs
RED = RGBColor(0xC0, 0x39, 0x2B)   # alerte (dépassement de seuil, anomalie ouverte)

NAVY_HEX = "052D62"
GREEN_HEX = "169D53"
LIGHT_HEX = "F1F5F9"
WHITE_HEX = "FFFFFF"

# Polices : titres géométriques façon maquette, corps humaniste lisible.
# « Segoe UI Semibold » est garantie sur Windows/Word et rend un sans-serif net
# proche de la maquette. Pour une fidélité au pixel au modèle (Montserrat),
# installer Montserrat sur le poste et remplacer HEADING_FONT par "Montserrat".
HEADING_FONT = "Segoe UI Semibold"
BODY_FONT = "Segoe UI"

CONTENT_WIDTH_CM = 16.0  # largeur utile avec marges de 2,5 cm sur A4
# Alignement des bandeaux : le fond coloré déborde de BANDEAU_PAD_CM à gauche (via
# un retrait de tableau négatif) ; la marge interne gauche des cellules
# (BANDEAU_TEXT_PAD_TW) est calibrée pour que le TEXTE des titres retombe pile sur
# la marge du corps (le +0,17 cm compense l'inset de cellule constaté au rendu).
BANDEAU_PAD_CM = 0.35
BANDEAU_TEXT_PAD_TW = int((BANDEAU_PAD_CM + 0.17) * 567)


# ─── Primitives XML bas niveau ────────────────────────────────────────────────

def _shade(parent_pr, fill_hex: str) -> None:
    """Applique un fond `fill_hex` à un pPr (paragraphe) ou tcPr (cellule)."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    parent_pr.append(shd)


def _bottom_border(paragraph, color_hex: str, size: int = 12) -> None:
    """Filet inférieur (soulignement de bloc) sous un paragraphe."""
    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color_hex)
    borders.append(bottom)
    p_pr.append(borders)


def _no_table_borders(table) -> None:
    """Supprime toutes les bordures d'un tableau (utilisé comme conteneur mis en fond)."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        borders.append(e)
    tbl_pr.append(borders)


def _set_table_width(table, cm: float) -> None:
    """Fixe la largeur d'un tableau (désactive l'ajustement auto) pour que tous
    les blocs pleine largeur aient EXACTEMENT la même longueur d'une page à l'autre."""
    table.autofit = False
    table.allow_autofit = False
    tbl_pr = table._tbl.tblPr
    existing = tbl_pr.find(qn("w:tblW"))
    if existing is not None:
        tbl_pr.remove(existing)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int(cm * 567)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_w)


def _bleed_left(table, pad_cm: float = BANDEAU_PAD_CM) -> None:
    """Retrait de tableau négatif : le bord gauche du bandeau déborde de `pad_cm`
    dans la marge, de sorte que le texte (décalé de la marge interne) retombe
    pile sur la marge du corps. Rend l'alignement identique dans Word et LO."""
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblInd"))
    if old is not None:
        tbl_pr.remove(old)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), str(-int(pad_cm * 567)))
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)


def _cell_shade(cell, fill_hex: str) -> None:
    _shade(cell._tc.get_or_add_tcPr(), fill_hex)


def _add_page_background(paragraph, w_pt: float, h_pt: float, fill_hex: str) -> None:
    """Ajoute un rectangle plein ancré à la page (derrière le texte) via VML.

    Étant ancré et « behindDoc », il n'occupe pas de place dans le flux : il ne
    peut donc pas provoquer de page blanche (contrairement à un tableau pleine
    hauteur dont Word pagine le paragraphe suiveur de façon imprévisible)."""
    run = paragraph.add_run()
    xml = (
        '<w:pict '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml">'
        f'<v:rect style="position:absolute;left:0;top:0;'
        f'width:{w_pt:.1f}pt;height:{h_pt:.1f}pt;'
        f'mso-position-horizontal:left;mso-position-horizontal-relative:page;'
        f'mso-position-vertical:top;mso-position-vertical-relative:page;'
        f'z-index:-251658240" fillcolor="#{fill_hex}" stroked="f"/>'
        f'</w:pict>'
    )
    run._r.append(parse_xml(xml))


def _cell_margins(cell, top=140, bottom=140, left=220, right=220) -> None:
    """Marges internes d'une cellule (en twips)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tc_pr.append(m)


def _set_row_exact_height(row, cm: float) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(cm * 567)))  # 1 cm = 567 twips
    h.set(qn("w:hRule"), "exact")
    tr_pr.append(h)


# ─── Styles de base ───────────────────────────────────────────────────────────

def apply_base_styles(doc) -> None:
    """Configure les styles Normal / Titres selon la charte."""
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.space_after = Pt(8)
    pf.line_spacing = 1.25

    for lvl, size in ((1, 15), (2, 12), (3, 11)):
        try:
            h = doc.styles[f"Heading {lvl}"]
        except KeyError:
            continue
        h.font.name = HEADING_FONT
        h.font.size = Pt(size)
        h.font.bold = True
        h.font.color.rgb = NAVY


def set_margins(section, top=2.5, bottom=2.2, left=2.5, right=2.5) -> None:
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


# ─── Pied de page numéroté ────────────────────────────────────────────────────

def add_page_number(section) -> None:
    """Numéro de page en bas à droite (façon maquette).

    Le champ PAGE embarque un fldChar « separate » + une valeur en cache (« 1 »)
    pour que le numéro s'affiche même sans recalcul des champs (conversion
    LibreOffice, aperçu) ; Word le met à jour à l'ouverture.
    """
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.text = ""
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    def _run():
        r = p.add_run()
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = NAVY
        r.font.name = HEADING_FONT
        return r

    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); _run()._r.append(fb)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "
    _run()._r.append(it)
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate"); _run()._r.append(fs)
    _run().text = "1"  # valeur en cache
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); _run()._r.append(fe)


# ─── Page de garde ────────────────────────────────────────────────────────────

def cover_page(doc, titre: str, sous_titre: str = "", meta: list[str] | None = None,
               cabinet: dict | None = None) -> None:
    """Page de garde pleine page bleu marine, titre blanc géant, filet vert.

    Le fond marine est un RECTANGLE ANCRÉ à la page (VML, derrière le texte) :
    il ne consomme pas d'espace dans le flux, donc il ne peut pas générer de page
    blanche en Word. Le texte de garde est posé en paragraphes normaux ; le corps
    démarre en page 2 via un saut NEW_PAGE (pieds de page indépendants).
    """
    cabinet = cabinet or {}
    meta = meta or []

    sec = doc.sections[0]
    # Marges de garde : texte inséré, titre positionné dans le tiers supérieur.
    set_margins(sec, top=6.0, bottom=2.0, left=2.5, right=2.5)
    w_pt = sec.page_width.pt
    h_pt = sec.page_height.pt

    # Titre géant (multi-lignes via '\n') — porte aussi le rectangle de fond.
    p_titre = doc.add_paragraph()
    p_titre.paragraph_format.space_after = Pt(6)
    _add_page_background(p_titre, w_pt, h_pt, NAVY_HEX)
    for i, ligne in enumerate(titre.split("\n")):
        if i:
            p_titre.add_run().add_break()
        r = p_titre.add_run(ligne)
        r.bold = True
        r.font.size = Pt(44)
        r.font.name = HEADING_FONT
        r.font.color.rgb = WHITE

    if sous_titre:
        p_sub = doc.add_paragraph()
        r = p_sub.add_run(sous_titre)
        r.font.size = Pt(22)
        r.font.name = HEADING_FONT
        r.font.color.rgb = WHITE

    # Filet vert horizontal.
    bar = doc.add_paragraph()
    bar.paragraph_format.space_before = Pt(10)
    _green_bar_run(bar, width_chars=28)

    # Méta (entité, référentiels…).
    for txt in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(txt)
        r.font.size = Pt(11)
        r.font.color.rgb = WHITE

    # Identité cabinet, un peu plus bas.
    if cabinet.get("nom"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(40)
        r = p.add_run(cabinet["nom"])
        r.bold = True
        r.font.size = Pt(14)
        r.font.name = HEADING_FONT
        r.font.color.rgb = WHITE
        sous = []
        if cabinet.get("forme_juridique"):
            sous.append(cabinet["forme_juridique"])
        if cabinet.get("adresse_ville"):
            sous.append(cabinet["adresse_ville"])
        if sous:
            p2 = doc.add_paragraph()
            r2 = p2.add_run("  ·  ".join(sous))
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0xC7, 0xD2, 0xE0)

    # Corps en page 2 (saut NEW_PAGE → pieds de page indépendants).
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    set_margins(body)
    add_page_number(body)


def _green_bar_run(paragraph, width_chars: int = 20) -> None:
    """Filet vert épais rendu par un fond vert sur des espaces insécables."""
    p_pr = paragraph._p.get_or_add_pPr()
    _shade(p_pr, GREEN_HEX)
    # hauteur du filet via bordures identiques vertes
    run = paragraph.add_run(" " * width_chars)
    run.font.size = Pt(4)


# ─── Bandeaux de section ──────────────────────────────────────────────────────

def section_header(doc, titre: str, numero: str | None = None) -> None:
    """Bandeau de section marine pleine largeur, texte blanc (façon « pilule »)."""
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    table = doc.add_table(rows=1, cols=1)
    _no_table_borders(table)
    _set_table_width(table, CONTENT_WIDTH_CM + BANDEAU_PAD_CM)
    _bleed_left(table)
    cell = table.cell(0, 0)
    cell.width = Cm(CONTENT_WIDTH_CM + BANDEAU_PAD_CM)
    _cell_shade(cell, NAVY_HEX)
    _cell_margins(cell, top=200, bottom=200, left=BANDEAU_TEXT_PAD_TW, right=360)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    libelle = f"{numero}. {titre}" if numero else titre
    r = p.add_run(libelle.upper())
    r.bold = True
    r.font.size = Pt(17)
    r.font.name = HEADING_FONT
    r.font.color.rgb = WHITE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def sous_titre(doc, texte: str):
    """Sous-titre marine gras souligné de vert."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(texte)
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = HEADING_FONT
    r.font.color.rgb = NAVY
    _bottom_border(p, GREEN_HEX, size=18)
    return p


def bande_verte(doc, texte: str):
    """Bandeau vert plein, texte blanc (mise en exergue d'un intitulé)."""
    table = doc.add_table(rows=1, cols=1)
    _no_table_borders(table)
    _set_table_width(table, CONTENT_WIDTH_CM + BANDEAU_PAD_CM)
    _bleed_left(table)
    cell = table.cell(0, 0)
    cell.width = Cm(CONTENT_WIDTH_CM + BANDEAU_PAD_CM)
    _cell_shade(cell, GREEN_HEX)
    _cell_margins(cell, top=120, bottom=120, left=BANDEAU_TEXT_PAD_TW, right=260)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(texte)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = HEADING_FONT
    r.font.color.rgb = WHITE
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def divider(doc) -> None:
    """Court filet vert de séparation."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    _green_bar_run(p, width_chars=16)


def para(doc, texte: str, *, justify: bool = True, bold: bool = False,
         italic: bool = False, size: float = 10.5, color: RGBColor | None = None):
    """Paragraphe de corps, justifié par défaut (rendu « bloc » du modèle)."""
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texte)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    return p


def lead(doc, texte: str):
    """Paragraphe d'accroche en gras marine (chapô)."""
    return para(doc, texte, bold=True, color=NAVY)


def citation(doc, texte: str, auteur: str = "", fonction: str = "") -> None:
    """Encadré de citation marine, guillemets verts, texte blanc italique."""
    table = doc.add_table(rows=1, cols=1)
    _no_table_borders(table)
    _set_table_width(table, CONTENT_WIDTH_CM)
    cell = table.cell(0, 0)
    cell.width = Cm(CONTENT_WIDTH_CM)
    _cell_shade(cell, NAVY_HEX)
    _cell_margins(cell, top=300, bottom=300, left=420, right=420)

    q = cell.paragraphs[0]
    q.paragraph_format.space_after = Pt(2)
    rq = q.add_run("“")
    rq.bold = True
    rq.font.size = Pt(40)
    rq.font.name = HEADING_FONT
    rq.font.color.rgb = GREEN

    pt = cell.add_paragraph()
    rt = pt.add_run(texte)
    rt.italic = True
    rt.font.size = Pt(11.5)
    rt.font.color.rgb = WHITE

    if auteur or fonction:
        pa = cell.add_paragraph()
        pa.paragraph_format.space_before = Pt(8)
        if auteur:
            ra = pa.add_run(auteur)
            ra.bold = True
            ra.font.color.rgb = WHITE
            ra.font.size = Pt(10.5)
        if fonction:
            pa.add_run("\n")
            rf = pa.add_run(fonction)
            rf.font.color.rgb = RGBColor(0xC7, 0xD2, 0xE0)
            rf.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def sommaire(doc, entrees: list[tuple[str, str]]) -> None:
    """Sommaire façon maquette : intitulé marine gras + description, filet marine.

    `entrees` : liste de (titre_section, description courte).
    """
    section_header(doc, "Sommaire")
    for i, (titre, desc) in enumerate(entrees, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(f"{titre.upper()}")
        r.bold = True
        r.font.size = Pt(13)
        r.font.name = HEADING_FONT
        r.font.color.rgb = NAVY
        if desc:
            d = doc.add_paragraph()
            d.paragraph_format.left_indent = Cm(0.5)
            d.paragraph_format.space_after = Pt(6)
            rd = d.add_run(desc)
            rd.font.size = Pt(9.5)
            rd.font.color.rgb = MUTED
        _bottom_border(p, NAVY_HEX, size=6)
    # Le sommaire occupe sa propre page : le contenu démarre sur la page suivante.
    doc.add_page_break()


def info_table(doc, lignes: list[tuple[str, str]]) -> None:
    """Petit tableau clé/valeur discret (identité, chiffres-clés)."""
    if not lignes:
        return
    table = doc.add_table(rows=len(lignes), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _no_table_borders(table)
    _set_table_width(table, CONTENT_WIDTH_CM + BANDEAU_PAD_CM)
    _bleed_left(table)
    for i, (cle, val) in enumerate(lignes):
        c0, c1 = table.rows[i].cells
        c0.width = Cm(5.5)
        c1.width = Cm(CONTENT_WIDTH_CM + BANDEAU_PAD_CM - 5.5)
        _cell_shade(c0, LIGHT_HEX)
        _cell_margins(c0, top=60, bottom=60, left=BANDEAU_TEXT_PAD_TW, right=160)
        _cell_margins(c1, top=60, bottom=60, left=160, right=160)
        r0 = c0.paragraphs[0].add_run(cle)
        r0.bold = True
        r0.font.size = Pt(9.5)
        r0.font.color.rgb = NAVY
        r1 = c1.paragraphs[0].add_run(str(val))
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = INK
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
